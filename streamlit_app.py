#!/usr/bin/env python3
"""PDF 편집 및 변환 도구 — Streamlit 웹 앱"""

import streamlit as st
import fitz
import io
import os
import json
import math
from PIL import Image

try:
    from streamlit_drawable_canvas import st_canvas
    HAS_CANVAS = True
except ImportError:
    HAS_CANVAS = False

try:
    from streamlit_image_coordinates import streamlit_image_coordinates
    HAS_IMG_COORDS = True
except ImportError:
    HAS_IMG_COORDS = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# ─── 페이지 설정 ──────────────────────────────────────────────────
st.set_page_config(
    page_title="PDF 편집 도구",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ─── 스타일 ───────────────────────────────────────────────────────
st.markdown("""
<style>
[data-testid="stSidebar"] { background: #181825; }
[data-testid="stSidebar"] * { color: #cdd6f4 !important; }
.block-container { padding-top: 1rem; padding-bottom: 1rem; }
.stTabs [data-baseweb="tab-list"] { gap: 4px; }
.stTabs [data-baseweb="tab"] {
    background: #313244; color: #cdd6f4;
    border-radius: 6px 6px 0 0; padding: 6px 18px;
}
.stTabs [aria-selected="true"] {
    background: #89b4fa !important; color: #1e1e2e !important;
}
h1,h2,h3 { color: #89b4fa; }
.stButton>button {
    background: #313244; color: #cdd6f4;
    border: 1px solid #45475a; border-radius: 6px;
}
.stButton>button:hover { background: #45475a; }
.stDownloadButton>button {
    background: #a6e3a1 !important; color: #1e1e2e !important;
    font-weight: bold;
}
</style>
""", unsafe_allow_html=True)

# ─── 세션 상태 초기화 ─────────────────────────────────────────────
def init_state():
    defaults = {
        "pdf_bytes": None,
        "filename": "",
        "current_page": 0,
        "total_pages": 0,
        "zoom": 1.0,
        "undo_stack": [],
        "draw_tool": "freedraw",
        "draw_color": "#ff0000",
        "draw_size": 3,
        "text_size": 14,
        "canvas_key": 0,
        "text_click_pos": None,   # (x, y) in PDF coords
        "text_tool_active": False,
        "text_edit_tool_active": False,
        "selected_block": None,   # {"text", "rect", "fontsize", "color"}
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

init_state()

# ─── 헬퍼 함수 ───────────────────────────────────────────────────

def get_pdf():
    """세션에서 fitz.Document 반환"""
    if st.session_state.pdf_bytes:
        return fitz.open("pdf", st.session_state.pdf_bytes)
    return None

def save_pdf(doc):
    """fitz.Document → 세션 바이트"""
    buf = io.BytesIO()
    doc.save(buf, garbage=4, deflate=True)
    st.session_state.pdf_bytes = buf.getvalue()

def push_undo():
    if st.session_state.pdf_bytes:
        st.session_state.undo_stack.append(st.session_state.pdf_bytes)
        if len(st.session_state.undo_stack) > 20:
            st.session_state.undo_stack.pop(0)

def render_page(doc, page_num, width=800):
    """페이지를 PIL Image로 렌더"""
    page = doc[page_num]
    scale = width / page.rect.width
    mat = fitz.Matrix(scale, scale)
    pix = page.get_pixmap(matrix=mat, alpha=False)
    return Image.open(io.BytesIO(pix.tobytes("png"))), scale

def hex_to_rgb01(hex_color):
    h = hex_color.lstrip("#")
    r, g, b = int(h[0:2],16)/255, int(h[2:4],16)/255, int(h[4:6],16)/255
    return r, g, b

def canvas_to_pdf(cx, cy, scale):
    return cx / scale, cy / scale

def find_text_block_at(page, pdf_x, pdf_y, tolerance=8):
    """클릭 위치에서 가장 가까운 텍스트 블록 반환"""
    pt = fitz.Point(pdf_x, pdf_y)
    best = None
    best_area = float("inf")
    for block in page.get_text("dict")["blocks"]:
        if block.get("type") != 0:
            continue
        rect = fitz.Rect(block["bbox"])
        expanded = rect + (-tolerance, -tolerance, tolerance, tolerance)
        if expanded.contains(pt):
            area = rect.width * rect.height
            if area < best_area:
                best_area = area
                # 첫 번째 span에서 폰트 정보 추출
                spans = [s for ln in block.get("lines", []) for s in ln.get("spans", [])]
                text = "".join(s["text"] for s in spans)
                fontsize = spans[0]["size"] if spans else 12
                color_int = spans[0].get("color", 0) if spans else 0
                r = ((color_int >> 16) & 0xFF) / 255
                g = ((color_int >> 8)  & 0xFF) / 255
                b = ( color_int        & 0xFF) / 255
                best = {
                    "text": text,
                    "rect": rect,
                    "fontsize": round(fontsize, 1),
                    "color": (r, g, b),
                    "color_hex": "#{:02x}{:02x}{:02x}".format(
                        int(r*255), int(g*255), int(b*255)),
                }
    return best

# ─── 사이드바 ─────────────────────────────────────────────────────
with st.sidebar:
    st.markdown("## 📄 PDF 도구")
    st.divider()

    # 파일 업로드
    uploaded = st.file_uploader("PDF 파일 열기", type=["pdf"],
                                  label_visibility="collapsed")
    if uploaded:
        data = uploaded.read()
        if data != st.session_state.pdf_bytes:
            st.session_state.pdf_bytes = data
            st.session_state.filename = uploaded.name
            doc = fitz.open("pdf", data)
            st.session_state.total_pages = len(doc)
            st.session_state.current_page = 0
            st.session_state.undo_stack = []
            st.session_state.canvas_key += 1
            doc.close()
            st.rerun()

    if st.session_state.pdf_bytes:
        st.success(f"📂 {st.session_state.filename}")
        st.caption(f"{st.session_state.total_pages}페이지")
        st.divider()

        # 페이지 네비게이션
        st.markdown("**페이지 이동**")
        col1, col2, col3 = st.columns([1,2,1])
        with col1:
            if st.button("◀", use_container_width=True):
                if st.session_state.current_page > 0:
                    st.session_state.current_page -= 1
                    st.session_state.canvas_key += 1
                    st.rerun()
        with col2:
            page_input = st.number_input("", min_value=1,
                max_value=st.session_state.total_pages,
                value=st.session_state.current_page + 1,
                label_visibility="collapsed")
            if page_input - 1 != st.session_state.current_page:
                st.session_state.current_page = page_input - 1
                st.session_state.canvas_key += 1
                st.rerun()
        with col3:
            if st.button("▶", use_container_width=True):
                if st.session_state.current_page < st.session_state.total_pages - 1:
                    st.session_state.current_page += 1
                    st.session_state.canvas_key += 1
                    st.rerun()

        st.caption(f"{st.session_state.current_page+1} / {st.session_state.total_pages}")
        st.divider()

        # 확대
        st.markdown("**확대/축소**")
        zoom_val = st.slider("", 0.5, 3.0,
                             st.session_state.zoom, 0.25,
                             label_visibility="collapsed")
        if zoom_val != st.session_state.zoom:
            st.session_state.zoom = zoom_val
            st.session_state.canvas_key += 1
            st.rerun()

        st.divider()

        # 저장
        st.markdown("**저장**")
        fname = st.session_state.filename or "output.pdf"
        st.download_button("⬇ PDF 다운로드", st.session_state.pdf_bytes,
                           file_name=fname, mime="application/pdf",
                           use_container_width=True)

        if st.button("↩ 실행취소", use_container_width=True):
            if st.session_state.undo_stack:
                st.session_state.pdf_bytes = st.session_state.undo_stack.pop()
                doc = fitz.open("pdf", st.session_state.pdf_bytes)
                st.session_state.total_pages = len(doc)
                doc.close()
                st.session_state.canvas_key += 1
                st.rerun()
            else:
                st.warning("되돌릴 내용이 없습니다")

# ─── 메인 영역 ────────────────────────────────────────────────────
if not st.session_state.pdf_bytes:
    st.markdown("""
    <div style="text-align:center; padding:80px; color:#585b70;">
        <div style="font-size:72px">📄</div>
        <h2 style="color:#89b4fa">PDF 편집 및 변환 도구</h2>
        <p style="font-size:16px">왼쪽 사이드바에서 PDF 파일을 업로드하세요</p>
        <br>
        <p>✏️ 뷰어에서 바로 텍스트·도형·필기 추가</p>
        <p>🔄 이미지, Word, 텍스트로 변환</p>
        <p>📎 PDF 병합 / 페이지 분할</p>
        <p>💧 워터마크 · 🔒 비밀번호 보호</p>
    </div>
    """, unsafe_allow_html=True)
    st.stop()

# ─── 탭 ──────────────────────────────────────────────────────────
tab_viewer, tab_edit, tab_convert, tab_merge, tab_watermark = st.tabs([
    "  🖊️ 뷰어 / 인라인 편집  ",
    "  ✂️ 페이지 편집  ",
    "  🔄 변환  ",
    "  📎 병합 / 분할  ",
    "  💧 워터마크 / 보안  ",
])

# ════════════════════════════════════════════════
# 탭 1: 뷰어 + 인라인 편집
# ════════════════════════════════════════════════
with tab_viewer:
    doc = get_pdf()
    if not doc:
        st.stop()

    display_width = int(800 * st.session_state.zoom)
    img, scale = render_page(doc, st.session_state.current_page, display_width)
    doc.close()
    canvas_h = img.height

    # ── 도구 선택 ──
    st.markdown("##### 편집 도구")
    tcol = st.columns(11)
    tool_labels = {
        "freedraw": "✏️ 펜", "line": "— 선", "rect": "▭ 사각형",
        "circle": "◯ 원", "transform": "↖ 선택",
    }

    def _is_draw_active(t):
        return not st.session_state.text_tool_active and not st.session_state.text_edit_tool_active and st.session_state.draw_tool == t

    with tcol[0]:
        if st.button("✏️ 펜", use_container_width=True,
                     type="primary" if _is_draw_active("freedraw") else "secondary"):
            st.session_state.draw_tool = "freedraw"
            st.session_state.text_tool_active = False
            st.session_state.text_edit_tool_active = False
            st.rerun()
    with tcol[1]:
        if st.button("— 선", use_container_width=True,
                     type="primary" if _is_draw_active("line") else "secondary"):
            st.session_state.draw_tool = "line"
            st.session_state.text_tool_active = False
            st.session_state.text_edit_tool_active = False
            st.rerun()
    with tcol[2]:
        if st.button("▭ 사각형", use_container_width=True,
                     type="primary" if _is_draw_active("rect") else "secondary"):
            st.session_state.draw_tool = "rect"
            st.session_state.text_tool_active = False
            st.session_state.text_edit_tool_active = False
            st.rerun()
    with tcol[3]:
        if st.button("◯ 원", use_container_width=True,
                     type="primary" if _is_draw_active("circle") else "secondary"):
            st.session_state.draw_tool = "circle"
            st.session_state.text_tool_active = False
            st.session_state.text_edit_tool_active = False
            st.rerun()
    with tcol[4]:
        if st.button("↖ 선택", use_container_width=True,
                     type="primary" if _is_draw_active("transform") else "secondary"):
            st.session_state.draw_tool = "transform"
            st.session_state.text_tool_active = False
            st.session_state.text_edit_tool_active = False
            st.rerun()
    with tcol[5]:
        if st.button("🔤 텍스트 추가", use_container_width=True,
                     type="primary" if st.session_state.text_tool_active else "secondary"):
            st.session_state.text_tool_active = True
            st.session_state.text_edit_tool_active = False
            st.session_state.text_click_pos = None
            st.rerun()
    with tcol[6]:
        if st.button("✏️ 텍스트 편집", use_container_width=True,
                     type="primary" if st.session_state.text_edit_tool_active else "secondary"):
            st.session_state.text_edit_tool_active = True
            st.session_state.text_tool_active = False
            st.session_state.selected_block = None
            st.rerun()

    with tcol[7]:
        st.session_state.draw_color = st.color_picker(
            "색상", st.session_state.draw_color, label_visibility="collapsed")
    with tcol[8]:
        st.session_state.draw_size = st.slider(
            "굵기", 1, 15, st.session_state.draw_size,
            label_visibility="collapsed")
    with tcol[9]:
        fill_shape = st.checkbox("채우기", False)

    cur_tool = st.session_state.draw_tool

    if st.session_state.text_edit_tool_active:
        st.caption("현재 도구: **✏️ 텍스트 편집**  |  수정할 텍스트를 PDF 위에서 클릭하세요")
    elif st.session_state.text_tool_active:
        st.caption("현재 도구: **🔤 텍스트 추가**  |  텍스트를 추가할 위치를 클릭하세요")
    else:
        st.caption(f"현재 도구: **{tool_labels.get(cur_tool, cur_tool)}**  |  색상: {st.session_state.draw_color}  |  굵기: {st.session_state.draw_size}")

    # ── 텍스트 도구: 클릭 위치 선택 ──
    if st.session_state.text_tool_active and HAS_IMG_COORDS:
        st.info("📍 텍스트를 추가할 위치를 PDF 위에서 클릭하세요")
        coords = streamlit_image_coordinates(img, key=f"img_coords_{st.session_state.canvas_key}")

        if coords is not None:
            # 클릭 좌표를 PDF 좌표로 변환
            pdf_x = coords["x"] / scale
            pdf_y = coords["y"] / scale
            st.session_state.text_click_pos = (pdf_x, pdf_y)

        if st.session_state.text_click_pos:
            px, py = st.session_state.text_click_pos
            st.success(f"📍 선택된 위치: X={px:.1f}, Y={py:.1f} (PDF 좌표)")
            ti_col1, ti_col2, ti_col3 = st.columns([4, 1, 1])
            with ti_col1:
                inline_text = st.text_input("추가할 텍스트", placeholder="텍스트를 입력하세요...",
                                             key="inline_text_input")
            with ti_col2:
                inline_size = st.number_input("크기", value=14, min_value=6, key="inline_text_size")
            with ti_col3:
                inline_color = st.color_picker("색상", "#000000", key="inline_text_color")

            btn_col1, btn_col2 = st.columns(2)
            with btn_col1:
                if st.button("✅ 텍스트 추가", type="primary", use_container_width=True):
                    if inline_text:
                        push_undo()
                        doc = get_pdf()
                        page = doc[st.session_state.current_page]
                        r, g, b = hex_to_rgb01(inline_color)
                        page.insert_text((px, py), inline_text,
                                          fontsize=inline_size, color=(r, g, b))
                        save_pdf(doc)
                        doc.close()
                        st.session_state.text_click_pos = None
                        st.session_state.canvas_key += 1
                        st.success(f'텍스트 추가됨: "{inline_text}"')
                        st.rerun()
                    else:
                        st.warning("텍스트를 입력하세요")
            with btn_col2:
                if st.button("❌ 취소", use_container_width=True):
                    st.session_state.text_click_pos = None
                    st.rerun()

    elif st.session_state.text_tool_active and not HAS_IMG_COORDS:
        # fallback: 수동 좌표 입력
        st.warning("streamlit-image-coordinates 패키지가 필요합니다: pip install streamlit-image-coordinates")
        tc1, tc2, tc3, tc4 = st.columns([3,1,1,1])
        with tc1:
            add_text_val = st.text_input("텍스트", placeholder="입력하세요...")
        with tc2:
            add_text_x = st.number_input("X", value=72, min_value=0)
        with tc3:
            add_text_y = st.number_input("Y", value=100, min_value=0)
        with tc4:
            add_text_size = st.number_input("크기", value=14, min_value=6)
        tcolor = st.color_picker("텍스트 색상", "#000000")
        if st.button("텍스트 PDF에 추가", type="primary"):
            if add_text_val:
                push_undo()
                doc = get_pdf()
                page = doc[st.session_state.current_page]
                r, g, b = hex_to_rgb01(tcolor)
                page.insert_text((add_text_x, add_text_y), add_text_val,
                                  fontsize=add_text_size, color=(r, g, b))
                save_pdf(doc)
                doc.close()
                st.session_state.canvas_key += 1
                st.success(f'텍스트 추가됨: "{add_text_val}"')
                st.rerun()

    # ── 텍스트 편집 도구 ──
    if st.session_state.text_edit_tool_active and HAS_IMG_COORDS:
        st.info("🖱️ 수정할 텍스트를 PDF 위에서 클릭하세요")
        edit_coords = streamlit_image_coordinates(img, key=f"edit_coords_{st.session_state.canvas_key}")

        if edit_coords is not None:
            pdf_ex = edit_coords["x"] / scale
            pdf_ey = edit_coords["y"] / scale
            doc = get_pdf()
            page = doc[st.session_state.current_page]
            blk = find_text_block_at(page, pdf_ex, pdf_ey)
            doc.close()
            if blk:
                st.session_state.selected_block = blk
            else:
                st.warning("해당 위치에서 텍스트를 찾지 못했습니다. 다시 클릭해보세요.")

        if st.session_state.selected_block:
            blk = st.session_state.selected_block
            st.success(f"선택된 텍스트: **{blk['text'][:60]}{'...' if len(blk['text'])>60 else ''}**")

            ed_col1, ed_col2, ed_col3 = st.columns([4, 1, 1])
            with ed_col1:
                edited_text = st.text_area("수정할 내용", value=blk["text"],
                                            height=80, key="edit_text_area")
            with ed_col2:
                edit_size = st.number_input("크기", value=float(blk["fontsize"]),
                                             min_value=4.0, step=0.5, key="edit_text_size")
            with ed_col3:
                edit_color = st.color_picker("색상", blk["color_hex"], key="edit_text_color")

            eb1, eb2 = st.columns(2)
            with eb1:
                if st.button("✅ 텍스트 교체", type="primary", use_container_width=True):
                    push_undo()
                    doc = get_pdf()
                    page = doc[st.session_state.current_page]
                    rect = blk["rect"]
                    # 기존 텍스트 영역을 흰색으로 가리기 (redaction)
                    page.add_redact_annot(rect, fill=(1, 1, 1))
                    page.apply_redactions()
                    # 새 텍스트 삽입 (baseline = rect.y1 기준)
                    r2, g2, b2 = hex_to_rgb01(edit_color)
                    page.insert_text(
                        (rect.x0, rect.y0 + edit_size),
                        edited_text,
                        fontsize=edit_size,
                        color=(r2, g2, b2),
                    )
                    save_pdf(doc)
                    doc.close()
                    st.session_state.selected_block = None
                    st.session_state.canvas_key += 1
                    st.success("텍스트가 교체됐습니다!")
                    st.rerun()
            with eb2:
                if st.button("❌ 취소", use_container_width=True, key="edit_cancel"):
                    st.session_state.selected_block = None
                    st.rerun()

    # ── 캔버스 (그리기 도구) ──
    if not st.session_state.text_tool_active and not st.session_state.text_edit_tool_active:
        if HAS_CANVAS:
            fill_color = (st.session_state.draw_color + "44") if fill_shape else "rgba(0,0,0,0)"

            canvas_result = st_canvas(
                fill_color=fill_color,
                stroke_width=st.session_state.draw_size,
                stroke_color=st.session_state.draw_color,
                background_image=img,
                update_streamlit=True,
                height=canvas_h,
                width=display_width,
                drawing_mode=cur_tool,
                key=f"canvas_{st.session_state.canvas_key}",
            )

            # 그리기 결과 PDF에 적용
            if st.button("🖊️ 그린 내용 PDF에 적용", type="primary"):
                if canvas_result.json_data and canvas_result.json_data.get("objects"):
                    objects = canvas_result.json_data["objects"]
                    if objects:
                        push_undo()
                        doc = get_pdf()
                        page = doc[st.session_state.current_page]
                        r, g, b = hex_to_rgb01(st.session_state.draw_color)
                        lw = max(st.session_state.draw_size * 0.5, 0.5)

                        for obj in objects:
                            otype = obj.get("type", "")

                            if otype == "path":
                                path_str = obj.get("path", [])
                                pts = []
                                for cmd in path_str:
                                    if cmd[0] in ("M", "L", "Q", "C") and len(cmd) >= 3:
                                        pts.append((cmd[-2] / scale, cmd[-1] / scale))
                                for i in range(len(pts) - 1):
                                    p1 = fitz.Point(pts[i])
                                    p2 = fitz.Point(pts[i+1])
                                    ann = page.add_line_annot(p1, p2)
                                    ann.set_colors(stroke=(r,g,b))
                                    ann.set_border(width=lw)
                                    ann.update()

                            elif otype == "line":
                                x1 = obj.get("x1", 0) + obj.get("left", 0)
                                y1 = obj.get("y1", 0) + obj.get("top", 0)
                                x2 = obj.get("x2", 0) + obj.get("left", 0)
                                y2 = obj.get("y2", 0) + obj.get("top", 0)
                                p1 = fitz.Point(x1 / scale, y1 / scale)
                                p2 = fitz.Point(x2 / scale, y2 / scale)
                                ann = page.add_line_annot(p1, p2)
                                ann.set_colors(stroke=(r,g,b))
                                ann.set_border(width=lw)
                                ann.update()

                            elif otype == "rect":
                                left   = obj.get("left", 0)
                                top    = obj.get("top", 0)
                                width  = obj.get("width", 0)  * obj.get("scaleX", 1)
                                height = obj.get("height", 0) * obj.get("scaleY", 1)
                                rect = fitz.Rect(
                                    left / scale, top / scale,
                                    (left + width) / scale, (top + height) / scale)
                                ann = page.add_rect_annot(rect)
                                ann.set_colors(stroke=(r,g,b),
                                               fill=(hex_to_rgb01(st.session_state.draw_color)
                                                     if fill_shape else None))
                                ann.set_border(width=lw)
                                ann.update()

                            elif otype == "ellipse":
                                left   = obj.get("left", 0)
                                top    = obj.get("top", 0)
                                rx     = obj.get("rx", 0) * obj.get("scaleX", 1)
                                ry     = obj.get("ry", 0) * obj.get("scaleY", 1)
                                cx_pdf = (left + rx) / scale
                                cy_pdf = (top + ry) / scale
                                rect = fitz.Rect(
                                    cx_pdf - rx/scale, cy_pdf - ry/scale,
                                    cx_pdf + rx/scale, cy_pdf + ry/scale)
                                ann = page.add_circle_annot(rect)
                                ann.set_colors(stroke=(r,g,b))
                                ann.set_border(width=lw)
                                ann.update()

                        save_pdf(doc)
                        doc.close()
                        st.session_state.canvas_key += 1
                        st.success(f"{len(objects)}개 객체가 PDF에 적용됐습니다!")
                        st.rerun()
                else:
                    st.info("캔버스에 그린 내용이 없습니다")
        else:
            st.image(img, use_container_width=False, width=display_width)
            st.warning("streamlit-drawable-canvas 설치가 필요합니다: pip install streamlit-drawable-canvas")

# ════════════════════════════════════════════════
# 탭 2: 페이지 편집
# ════════════════════════════════════════════════
with tab_edit:
    st.markdown("#### 페이지 관리")

    doc = get_pdf()
    if not doc:
        st.stop()
    total = len(doc)
    doc.close()

    col1, col2 = st.columns(2)

    with col1:
        st.markdown("**현재 페이지 작업**")
        cur = st.session_state.current_page + 1

        if st.button(f"🔄 페이지 {cur} 90° 회전", use_container_width=True):
            push_undo()
            doc = get_pdf()
            page = doc[st.session_state.current_page]
            page.set_rotation((page.rotation + 90) % 360)
            save_pdf(doc)
            doc.close()
            st.session_state.canvas_key += 1
            st.success("회전됨")
            st.rerun()

        if total > 1:
            if st.button(f"🗑️ 페이지 {cur} 삭제", use_container_width=True):
                push_undo()
                doc = get_pdf()
                doc.delete_page(st.session_state.current_page)
                save_pdf(doc)
                n = len(doc)
                doc.close()
                st.session_state.total_pages = n
                st.session_state.current_page = min(
                    st.session_state.current_page, n - 1)
                st.session_state.canvas_key += 1
                st.success("삭제됨")
                st.rerun()

        if st.button(f"📋 페이지 {cur} 복제", use_container_width=True):
            push_undo()
            doc = get_pdf()
            doc.copy_page(st.session_state.current_page,
                          st.session_state.current_page + 1)
            st.session_state.total_pages = len(doc)
            save_pdf(doc)
            doc.close()
            st.session_state.canvas_key += 1
            st.success("복제됨")
            st.rerun()

    with col2:
        st.markdown("**페이지 이동**")
        move_to = st.number_input("이동할 위치", 1, total,
                                   st.session_state.current_page + 1)
        if st.button("페이지 이동", use_container_width=True):
            push_undo()
            doc = get_pdf()
            doc.move_page(st.session_state.current_page, move_to - 1)
            save_pdf(doc)
            doc.close()
            st.session_state.current_page = move_to - 1
            st.session_state.canvas_key += 1
            st.success(f"페이지를 {move_to}번으로 이동")
            st.rerun()

    st.divider()

    # 이미지 삽입
    st.markdown("#### 이미지 삽입")
    ins_img = st.file_uploader("이미지 파일", type=["png","jpg","jpeg","bmp"],
                                key="ins_img_uploader")
    if ins_img:
        ic1, ic2, ic3, ic4 = st.columns(4)
        with ic1: ix = st.number_input("X", value=50, min_value=0)
        with ic2: iy = st.number_input("Y", value=50, min_value=0)
        with ic3: iw = st.number_input("너비", value=200, min_value=10)
        with ic4: ih = st.number_input("높이", value=200, min_value=10)

        if st.button("이미지 삽입", type="primary"):
            push_undo()
            doc = get_pdf()
            page = doc[st.session_state.current_page]
            img_bytes = ins_img.read()
            rect = fitz.Rect(ix, iy, ix + iw, iy + ih)
            page.insert_image(rect, stream=img_bytes)
            save_pdf(doc)
            doc.close()
            st.session_state.canvas_key += 1
            st.success("이미지 삽입됨")
            st.rerun()

    st.divider()

    # 텍스트 추출
    st.markdown("#### 텍스트 추출")
    exc1, exc2 = st.columns(2)
    with exc1:
        if st.button("현재 페이지 텍스트 추출", use_container_width=True):
            doc = get_pdf()
            text = doc[st.session_state.current_page].get_text()
            doc.close()
            st.session_state["extracted_text"] = text

    with exc2:
        if st.button("전체 텍스트 추출", use_container_width=True):
            doc = get_pdf()
            all_text = ""
            for i, page in enumerate(doc):
                all_text += f"=== 페이지 {i+1} ===\n{page.get_text()}\n\n"
            doc.close()
            st.session_state["extracted_text"] = all_text

    if "extracted_text" in st.session_state and st.session_state["extracted_text"]:
        et = st.session_state["extracted_text"]
        st.text_area("추출된 텍스트", et, height=200)
        st.download_button("📥 텍스트 파일로 저장", et.encode("utf-8"),
                           file_name="extracted.txt", mime="text/plain")

# ════════════════════════════════════════════════
# 탭 3: 변환
# ════════════════════════════════════════════════
with tab_convert:
    st.markdown("#### PDF 변환")

    c1, c2 = st.columns(2)

    with c1:
        # PDF → 이미지
        st.markdown("**📷 PDF → 이미지**")
        img_fmt = st.selectbox("형식", ["PNG", "JPEG", "BMP"])
        dpi_val = st.select_slider("DPI", [72, 96, 150, 200, 300, 600], 150)
        page_mode = st.radio("페이지", ["전체", "현재 페이지"], horizontal=True)

        if st.button("이미지로 변환", use_container_width=True):
            doc = get_pdf()
            mat = fitz.Matrix(dpi_val/72, dpi_val/72)
            pages = ([st.session_state.current_page]
                     if page_mode == "현재 페이지"
                     else range(len(doc)))

            if len(list(pages)) == 1:
                pages = ([st.session_state.current_page]
                         if page_mode == "현재 페이지"
                         else range(len(doc)))
                page = doc[list(pages)[0]]
                pix = page.get_pixmap(matrix=mat, alpha=False)
                img_bytes = pix.tobytes(img_fmt.lower())
                doc.close()
                st.download_button(
                    f"⬇ 이미지 다운로드 ({img_fmt})",
                    img_bytes,
                    file_name=f"page_{st.session_state.current_page+1}.{img_fmt.lower()}",
                    mime=f"image/{img_fmt.lower()}")
            else:
                import zipfile
                zip_buf = io.BytesIO()
                pages = range(len(doc))
                with zipfile.ZipFile(zip_buf, "w") as zf:
                    bar = st.progress(0)
                    for i, pn in enumerate(pages):
                        page = doc[pn]
                        pix = page.get_pixmap(matrix=mat, alpha=False)
                        img_b = pix.tobytes(img_fmt.lower())
                        zf.writestr(f"page_{pn+1}.{img_fmt.lower()}", img_b)
                        bar.progress((i+1)/len(doc))
                doc.close()
                st.download_button("⬇ ZIP 다운로드", zip_buf.getvalue(),
                                   file_name="pdf_images.zip",
                                   mime="application/zip")

        st.divider()

        # PDF → 텍스트
        st.markdown("**📝 PDF → 텍스트**")
        if st.button("텍스트 파일로 변환", use_container_width=True):
            doc = get_pdf()
            all_text = ""
            bar = st.progress(0)
            for i, page in enumerate(doc):
                all_text += f"=== 페이지 {i+1} ===\n{page.get_text()}\n\n"
                bar.progress((i+1)/len(doc))
            doc.close()
            st.download_button("⬇ TXT 다운로드", all_text.encode("utf-8"),
                               file_name="output.txt", mime="text/plain")

    with c2:
        # PDF → Word
        st.markdown("**📄 PDF → Word (.docx)**")
        if not HAS_DOCX:
            st.warning("python-docx 설치 필요: pip install python-docx")
        else:
            if st.button("Word로 변환", use_container_width=True):
                doc = get_pdf()
                word_doc = Document()
                fname_base = os.path.splitext(st.session_state.filename)[0]
                word_doc.add_heading(fname_base, 0)
                bar = st.progress(0)
                for i, page in enumerate(doc):
                    word_doc.add_heading(f"페이지 {i+1}", 2)
                    text = page.get_text()
                    word_doc.add_paragraph(text if text.strip() else "(텍스트 없음)")
                    bar.progress((i+1)/len(doc))
                doc.close()
                buf = io.BytesIO()
                word_doc.save(buf)
                st.download_button("⬇ DOCX 다운로드", buf.getvalue(),
                                   file_name=f"{fname_base}.docx",
                                   mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        st.divider()

        # 이미지 → PDF
        st.markdown("**🖼️ 이미지 → PDF**")
        uploaded_imgs = st.file_uploader(
            "이미지 파일들", type=["png","jpg","jpeg","bmp","tiff"],
            accept_multiple_files=True, key="imgs_to_pdf")
        if uploaded_imgs and st.button("PDF 생성", use_container_width=True):
            new_doc = fitz.open()
            bar = st.progress(0)
            for i, up_img in enumerate(uploaded_imgs):
                img = Image.open(up_img)
                w, h = img.size
                page = new_doc.new_page(width=w, height=h)
                img_bytes = io.BytesIO()
                img.save(img_bytes, format="PNG")
                page.insert_image(fitz.Rect(0, 0, w, h),
                                  stream=img_bytes.getvalue())
                bar.progress((i+1)/len(uploaded_imgs))
            buf = io.BytesIO()
            new_doc.save(buf, garbage=4, deflate=True)
            new_doc.close()
            st.download_button("⬇ PDF 다운로드", buf.getvalue(),
                               file_name="images_merged.pdf",
                               mime="application/pdf")

# ════════════════════════════════════════════════
# 탭 4: 병합 / 분할
# ════════════════════════════════════════════════
with tab_merge:
    col1, col2 = st.columns(2)

    with col1:
        st.markdown("#### 📎 PDF 병합")
        merge_files = st.file_uploader(
            "병합할 PDF 파일들 (순서대로 선택)",
            type=["pdf"], accept_multiple_files=True, key="merge_uploader")

        if merge_files and len(merge_files) >= 2:
            st.caption(f"{len(merge_files)}개 파일 선택됨:")
            for i, f in enumerate(merge_files):
                st.caption(f"  {i+1}. {f.name}")

            if st.button("병합하기", type="primary", use_container_width=True):
                merged = fitz.open()
                bar = st.progress(0)
                for i, f in enumerate(merge_files):
                    doc = fitz.open("pdf", f.read())
                    merged.insert_pdf(doc)
                    doc.close()
                    bar.progress((i+1)/len(merge_files))
                buf = io.BytesIO()
                merged.save(buf, garbage=4, deflate=True)
                merged.close()
                st.download_button("⬇ 병합된 PDF 다운로드", buf.getvalue(),
                                   file_name="merged.pdf",
                                   mime="application/pdf")
        elif merge_files:
            st.info("PDF를 2개 이상 선택하세요")

    with col2:
        st.markdown("#### ✂️ PDF 분할")
        doc = get_pdf()
        if not doc:
            doc.close()
            st.stop()
        total_p = len(doc)
        doc.close()

        split_mode = st.radio("분할 방식", ["페이지 범위", "각 페이지별", "N페이지씩"])

        if split_mode == "페이지 범위":
            range_str = st.text_input("범위 입력 (예: 1-3, 5, 7-10)")
            if st.button("분할하기", use_container_width=True) and range_str:
                pages = set()
                for part in range_str.split(","):
                    part = part.strip()
                    if "-" in part:
                        a, b = part.split("-", 1)
                        try:
                            for p in range(int(a.strip())-1, int(b.strip())):
                                pages.add(p)
                        except:
                            pass
                    else:
                        try:
                            pages.add(int(part)-1)
                        except:
                            pass
                pages = sorted(pages)
                if pages:
                    doc = get_pdf()
                    new_doc = fitz.open()
                    for p in pages:
                        if 0 <= p < total_p:
                            new_doc.insert_pdf(doc, from_page=p, to_page=p)
                    doc.close()
                    buf = io.BytesIO()
                    new_doc.save(buf, garbage=4, deflate=True)
                    new_doc.close()
                    st.download_button("⬇ 분할된 PDF 다운로드", buf.getvalue(),
                                       file_name="split_range.pdf",
                                       mime="application/pdf")

        elif split_mode == "각 페이지별":
            if st.button("각 페이지를 PDF로 분할", use_container_width=True):
                import zipfile
                zip_buf = io.BytesIO()
                doc = get_pdf()
                bar = st.progress(0)
                with zipfile.ZipFile(zip_buf, "w") as zf:
                    for i in range(total_p):
                        new_doc = fitz.open()
                        new_doc.insert_pdf(doc, from_page=i, to_page=i)
                        page_buf = io.BytesIO()
                        new_doc.save(page_buf)
                        new_doc.close()
                        zf.writestr(f"page_{i+1}.pdf", page_buf.getvalue())
                        bar.progress((i+1)/total_p)
                doc.close()
                st.download_button("⬇ ZIP 다운로드", zip_buf.getvalue(),
                                   file_name="split_pages.zip",
                                   mime="application/zip")

        else:
            n_pages = st.number_input("N (페이지 수)", 1, total_p, 1)
            if st.button(f"매 {n_pages}페이지씩 분할", use_container_width=True):
                import zipfile
                zip_buf = io.BytesIO()
                doc = get_pdf()
                count = 0
                with zipfile.ZipFile(zip_buf, "w") as zf:
                    for start in range(0, total_p, n_pages):
                        end = min(start + n_pages - 1, total_p - 1)
                        new_doc = fitz.open()
                        new_doc.insert_pdf(doc, from_page=start, to_page=end)
                        page_buf = io.BytesIO()
                        new_doc.save(page_buf)
                        new_doc.close()
                        zf.writestr(f"part_{count+1}.pdf", page_buf.getvalue())
                        count += 1
                doc.close()
                st.download_button("⬇ ZIP 다운로드", zip_buf.getvalue(),
                                   file_name="split_parts.zip",
                                   mime="application/zip")

# ════════════════════════════════════════════════
# 탭 5: 워터마크 / 보안
# ════════════════════════════════════════════════
with tab_watermark:
    wc1, wc2 = st.columns(2)

    with wc1:
        st.markdown("#### 💧 텍스트 워터마크")
        wm_text  = st.text_input("워터마크 텍스트", "CONFIDENTIAL")
        wm_size  = st.slider("폰트 크기", 10, 100, 40)
        wm_alpha = st.slider("투명도", 0.05, 1.0, 0.3, 0.05)
        wm_angle = st.slider("각도 (도)", 0, 360, 45)
        wm_color = st.color_picker("색상", "#888888")
        wm_pages = st.radio("적용 범위", ["전체 페이지", "현재 페이지"], horizontal=True)

        if st.button("워터마크 추가", type="primary", use_container_width=True):
            if wm_text:
                push_undo()
                doc = get_pdf()
                r, g, b = hex_to_rgb01(wm_color)
                pages = (range(len(doc)) if wm_pages == "전체 페이지"
                         else [st.session_state.current_page])
                for pn in pages:
                    page = doc[pn]
                    rect = page.rect
                    cx = rect.width / 2 - wm_size * len(wm_text) * 0.3
                    cy = rect.height / 2
                    page.insert_text(
                        (cx, cy), wm_text,
                        fontsize=wm_size,
                        color=(r, g, b),
                        rotate=wm_angle,
                        fill_opacity=wm_alpha,
                    )
                save_pdf(doc)
                doc.close()
                st.session_state.canvas_key += 1
                st.success("워터마크 추가됨!")
                st.rerun()

        st.divider()
        st.markdown("#### ⚙️ PDF 최적화")
        if st.button("파일 크기 최적화", use_container_width=True):
            orig_size = len(st.session_state.pdf_bytes)
            doc = get_pdf()
            buf = io.BytesIO()
            doc.save(buf, garbage=4, deflate=True,
                     clean=True, deflate_images=True, deflate_fonts=True)
            doc.close()
            new_bytes = buf.getvalue()
            new_size = len(new_bytes)
            saved = orig_size - new_size
            pct = saved / orig_size * 100 if orig_size > 0 else 0
            st.info(f"원본: {orig_size/1024:.1f} KB → 최적화: {new_size/1024:.1f} KB "
                    f"({pct:.1f}% 절약)")
            st.download_button("⬇ 최적화된 PDF 다운로드", new_bytes,
                               file_name="optimized.pdf",
                               mime="application/pdf")

    with wc2:
        st.markdown("#### 🔒 비밀번호 보호")
        pw = st.text_input("비밀번호", type="password")
        pw_confirm = st.text_input("비밀번호 확인", type="password")

        if st.button("비밀번호 설정 및 다운로드", type="primary",
                     use_container_width=True):
            if not pw:
                st.error("비밀번호를 입력하세요")
            elif pw != pw_confirm:
                st.error("비밀번호가 일치하지 않습니다")
            else:
                doc = get_pdf()
                buf = io.BytesIO()
                perm = (fitz.PDF_PERM_ACCESSIBILITY
                        | fitz.PDF_PERM_PRINT
                        | fitz.PDF_PERM_COPY
                        | fitz.PDF_PERM_ANNOTATE)
                doc.save(buf,
                         encryption=fitz.PDF_ENCRYPT_AES_256,
                         owner_pw=pw + "_owner",
                         user_pw=pw,
                         permissions=perm,
                         garbage=4, deflate=True)
                doc.close()
                st.download_button("⬇ 암호화된 PDF 다운로드", buf.getvalue(),
                                   file_name="encrypted.pdf",
                                   mime="application/pdf")
                st.success("AES-256으로 암호화됨")

        st.divider()
        st.markdown("#### 🔓 비밀번호 제거")
        rm_pw = st.text_input("현재 비밀번호", type="password", key="rm_pw")
        if st.button("비밀번호 제거", use_container_width=True):
            try:
                doc = fitz.open("pdf", st.session_state.pdf_bytes)
                if doc.is_encrypted:
                    if doc.authenticate(rm_pw):
                        buf = io.BytesIO()
                        doc.save(buf, garbage=4, deflate=True)
                        doc.close()
                        st.download_button("⬇ 잠금 해제된 PDF", buf.getvalue(),
                                           file_name="unlocked.pdf",
                                           mime="application/pdf")
                        st.success("비밀번호 제거됨")
                    else:
                        st.error("비밀번호가 올바르지 않습니다")
                        doc.close()
                else:
                    st.info("이 PDF는 암호화되어 있지 않습니다")
                    doc.close()
            except Exception as e:
                st.error(f"오류: {e}")
