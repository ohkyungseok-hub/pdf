#!/usr/bin/env python3
"""PDF 편집 및 변환 도구"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, simpledialog
import os
import sys
import threading
import io

try:
    import fitz  # PyMuPDF
except ImportError:
    print("PyMuPDF 설치 중...")
    os.system(f"{sys.executable} -m pip install PyMuPDF")
    import fitz

try:
    from PIL import Image, ImageTk
except ImportError:
    os.system(f"{sys.executable} -m pip install Pillow")
    from PIL import Image, ImageTk

try:
    from docx import Document
except ImportError:
    os.system(f"{sys.executable} -m pip install python-docx")
    from docx import Document

try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4, letter
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Paragraph
    from reportlab.lib.styles import getSampleStyleSheet
except ImportError:
    os.system(f"{sys.executable} -m pip install reportlab")
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4, letter


class PDFTool:
    def __init__(self, root):
        self.root = root
        self.root.title("PDF 편집 및 변환 도구")
        self.root.geometry("1200x800")
        self.root.configure(bg="#1e1e2e")

        self.current_pdf = None
        self.current_path = None
        self.current_page = 0
        self.total_pages = 0
        self.zoom = 1.0
        self.photo_images = []
        self.selected_pages = set()

        # 인라인 편집 상태
        self.edit_tool = "select"
        self.tool_btns = {}
        self.draw_start = None
        self.draw_preview = None
        self.freehand_points = []
        self.freehand_canvas_ids = []
        self.inline_entry_win = None
        self.inline_entry_pos = None
        self.draw_color = (1, 0, 0)          # PDF RGB (0-1)
        self.draw_color_hex = "#ff0000"
        self.draw_size = 2
        self.text_size = 12
        self.undo_stack = []

        self._setup_styles()
        self._build_ui()

    def _setup_styles(self):
        style = ttk.Style()
        style.theme_use("clam")
        style.configure("TNotebook", background="#1e1e2e", borderwidth=0)
        style.configure("TNotebook.Tab",
                        background="#313244", foreground="#cdd6f4",
                        padding=[12, 6], font=("Helvetica", 11))
        style.map("TNotebook.Tab",
                  background=[("selected", "#89b4fa")],
                  foreground=[("selected", "#1e1e2e")])
        style.configure("TFrame", background="#1e1e2e")
        style.configure("TLabel", background="#1e1e2e", foreground="#cdd6f4")
        style.configure("TButton",
                        background="#313244", foreground="#cdd6f4",
                        font=("Helvetica", 10), padding=[8, 4])
        style.map("TButton",
                  background=[("active", "#89b4fa")],
                  foreground=[("active", "#1e1e2e")])
        style.configure("Accent.TButton",
                        background="#89b4fa", foreground="#1e1e2e",
                        font=("Helvetica", 10, "bold"), padding=[10, 6])
        style.map("Accent.TButton",
                  background=[("active", "#74c7ec")])

    def _build_ui(self):
        # 상단 툴바
        toolbar = tk.Frame(self.root, bg="#181825", pady=8)
        toolbar.pack(fill="x")

        tk.Label(toolbar, text="  PDF 도구", bg="#181825",
                 fg="#89b4fa", font=("Helvetica", 14, "bold")).pack(side="left")

        btn_frame = tk.Frame(toolbar, bg="#181825")
        btn_frame.pack(side="left", padx=20)

        buttons = [
            ("열기", self.open_pdf, "#89b4fa"),
            ("저장", self.save_pdf, "#a6e3a1"),
            ("다른 이름으로 저장", self.save_pdf_as, "#a6e3a1"),
        ]
        for text, cmd, color in buttons:
            tk.Button(btn_frame, text=text, command=cmd,
                      bg=color, fg="#1e1e2e",
                      font=("Helvetica", 10, "bold"),
                      relief="flat", padx=10, pady=4,
                      cursor="hand2").pack(side="left", padx=3)

        # 상태바
        self.status_var = tk.StringVar(value="PDF 파일을 열어주세요")
        status_bar = tk.Label(self.root, textvariable=self.status_var,
                              bg="#11111b", fg="#a6adc8",
                              font=("Helvetica", 9), anchor="w", padx=10)
        status_bar.pack(side="bottom", fill="x")

        # 메인 영역
        main_paned = tk.PanedWindow(self.root, orient="horizontal",
                                    bg="#1e1e2e", sashwidth=4,
                                    sashrelief="flat")
        main_paned.pack(fill="both", expand=True, padx=0, pady=0)

        # 왼쪽: 페이지 썸네일 패널
        left_frame = tk.Frame(main_paned, bg="#181825", width=180)
        main_paned.add(left_frame, minsize=150)

        tk.Label(left_frame, text="페이지", bg="#181825",
                 fg="#89b4fa", font=("Helvetica", 11, "bold")).pack(pady=8)

        thumb_scroll = tk.Scrollbar(left_frame, orient="vertical")
        self.thumb_canvas = tk.Canvas(left_frame, bg="#181825",
                                       yscrollcommand=thumb_scroll.set,
                                       highlightthickness=0)
        thumb_scroll.config(command=self.thumb_canvas.yview)
        thumb_scroll.pack(side="right", fill="y")
        self.thumb_canvas.pack(fill="both", expand=True)
        self.thumb_inner = tk.Frame(self.thumb_canvas, bg="#181825")
        self.thumb_canvas.create_window((0, 0), window=self.thumb_inner, anchor="nw")
        self.thumb_inner.bind("<Configure>",
            lambda e: self.thumb_canvas.configure(
                scrollregion=self.thumb_canvas.bbox("all")))

        # 오른쪽: 탭 패널
        right_frame = tk.Frame(main_paned, bg="#1e1e2e")
        main_paned.add(right_frame, minsize=600)

        self.notebook = ttk.Notebook(right_frame)
        self.notebook.pack(fill="both", expand=True, padx=5, pady=5)

        self._build_viewer_tab()
        self._build_edit_tab()
        self._build_convert_tab()
        self._build_merge_split_tab()
        self._build_watermark_tab()

    def _build_viewer_tab(self):
        tab = ttk.Frame(self.notebook)
        self.notebook.add(tab, text="  뷰어/편집  ")

        # ── 상단: 페이지 네비게이션 바 ─────────────────────────
        nav = tk.Frame(tab, bg="#313244", pady=5)
        nav.pack(fill="x")

        tk.Button(nav, text="◀◀", command=self.first_page,
                  bg="#313244", fg="#cdd6f4", relief="flat",
                  font=("Helvetica", 12), cursor="hand2").pack(side="left", padx=2)
        tk.Button(nav, text="◀", command=self.prev_page,
                  bg="#313244", fg="#cdd6f4", relief="flat",
                  font=("Helvetica", 12), cursor="hand2").pack(side="left", padx=2)

        self.page_var = tk.StringVar(value="0 / 0")
        tk.Label(nav, textvariable=self.page_var,
                 bg="#313244", fg="#cdd6f4",
                 font=("Helvetica", 10), width=10).pack(side="left", padx=5)

        tk.Button(nav, text="▶", command=self.next_page,
                  bg="#313244", fg="#cdd6f4", relief="flat",
                  font=("Helvetica", 12), cursor="hand2").pack(side="left", padx=2)
        tk.Button(nav, text="▶▶", command=self.last_page,
                  bg="#313244", fg="#cdd6f4", relief="flat",
                  font=("Helvetica", 12), cursor="hand2").pack(side="left", padx=2)

        tk.Label(nav, text="  |", bg="#313244", fg="#45475a").pack(side="left")
        tk.Button(nav, text="−", command=self.zoom_out,
                  bg="#313244", fg="#cdd6f4", relief="flat",
                  font=("Helvetica", 14), cursor="hand2").pack(side="left", padx=2)
        self.zoom_var = tk.StringVar(value="100%")
        tk.Label(nav, textvariable=self.zoom_var,
                 bg="#313244", fg="#cdd6f4",
                 font=("Helvetica", 10), width=6).pack(side="left")
        tk.Button(nav, text="+", command=self.zoom_in,
                  bg="#313244", fg="#cdd6f4", relief="flat",
                  font=("Helvetica", 14), cursor="hand2").pack(side="left", padx=2)
        tk.Button(nav, text="맞춤", command=self.zoom_fit,
                  bg="#313244", fg="#cdd6f4", relief="flat",
                  font=("Helvetica", 10), cursor="hand2").pack(side="left", padx=3)

        # 실행취소
        tk.Label(nav, text="  |", bg="#313244", fg="#45475a").pack(side="left")
        tk.Button(nav, text="↩ 실행취소", command=self.undo,
                  bg="#313244", fg="#f38ba8", relief="flat",
                  font=("Helvetica", 10), cursor="hand2").pack(side="left", padx=5)

        # ── 편집 도구 바 ──────────────────────────────────────
        toolbar = tk.Frame(tab, bg="#1e1e2e", pady=4)
        toolbar.pack(fill="x", padx=4)

        tk.Label(toolbar, text="도구:", bg="#1e1e2e",
                 fg="#a6adc8", font=("Helvetica", 9)).pack(side="left", padx=(4, 2))

        tools = [
            ("select",    "↖ 선택",   "#6c7086"),
            ("text",      "T 텍스트",  "#89b4fa"),
            ("highlight", "▬ 형광펜", "#f9e2af"),
            ("rect",      "▭ 사각형",  "#a6e3a1"),
            ("pencil",    "✏ 펜",     "#cba6f7"),
            ("arrow",     "→ 화살표",  "#89dceb"),
            ("redact",    "█ 가리기",  "#f38ba8"),
        ]
        for tool_id, label, color in tools:
            btn = tk.Button(toolbar, text=label,
                            bg="#313244", fg=color,
                            relief="flat", padx=8, pady=3,
                            font=("Helvetica", 9, "bold"),
                            cursor="hand2",
                            command=lambda t=tool_id: self._set_tool(t))
            btn.pack(side="left", padx=2)
            self.tool_btns[tool_id] = btn

        tk.Label(toolbar, text=" | ", bg="#1e1e2e", fg="#45475a").pack(side="left")

        # 색상 선택
        tk.Label(toolbar, text="색상:", bg="#1e1e2e",
                 fg="#a6adc8", font=("Helvetica", 9)).pack(side="left")
        colors = [
            ("#ff0000", (1,0,0)),
            ("#0000ff", (0,0,1)),
            ("#00aa00", (0,0.67,0)),
            ("#ff8800", (1,0.53,0)),
            ("#aa00aa", (0.67,0,0.67)),
            ("#000000", (0,0,0)),
        ]
        for hex_c, rgb_c in colors:
            b = tk.Button(toolbar, bg=hex_c, width=2, height=1,
                          relief="ridge", cursor="hand2",
                          command=lambda h=hex_c, r=rgb_c: self._set_color(h, r))
            b.pack(side="left", padx=1)
        self.color_indicator = tk.Label(toolbar, bg=self.draw_color_hex,
                                         width=3, relief="sunken")
        self.color_indicator.pack(side="left", padx=3)

        tk.Label(toolbar, text="굵기:", bg="#1e1e2e",
                 fg="#a6adc8", font=("Helvetica", 9)).pack(side="left")
        self.size_var = tk.StringVar(value="2")
        size_cb = ttk.Combobox(toolbar, textvariable=self.size_var,
                                values=["1","2","3","5","8"],
                                width=3, state="readonly")
        size_cb.pack(side="left", padx=2)
        size_cb.bind("<<ComboboxSelected>>",
                     lambda e: setattr(self, "draw_size", int(self.size_var.get())))

        tk.Label(toolbar, text=" 텍스트크기:", bg="#1e1e2e",
                 fg="#a6adc8", font=("Helvetica", 9)).pack(side="left")
        self.tsize_var = tk.StringVar(value="12")
        tsize_cb = ttk.Combobox(toolbar, textvariable=self.tsize_var,
                                 values=["8","10","12","14","16","18","24","32"],
                                 width=4, state="readonly")
        tsize_cb.pack(side="left", padx=2)
        tsize_cb.bind("<<ComboboxSelected>>",
                      lambda e: setattr(self, "text_size", int(self.tsize_var.get())))

        # 도구 힌트
        self.tool_hint_var = tk.StringVar(value="↖ 선택 모드")
        tk.Label(toolbar, textvariable=self.tool_hint_var,
                 bg="#1e1e2e", fg="#585b70",
                 font=("Helvetica", 9, "italic")).pack(side="right", padx=8)

        # ── 뷰어 캔버스 ──────────────────────────────────────
        viewer_frame = tk.Frame(tab, bg="#1e1e2e")
        viewer_frame.pack(fill="both", expand=True)

        v_scroll = tk.Scrollbar(viewer_frame, orient="vertical", bg="#313244")
        h_scroll = tk.Scrollbar(viewer_frame, orient="horizontal", bg="#313244")

        self.viewer_canvas = tk.Canvas(viewer_frame,
                                        bg="#2a2a3e",
                                        yscrollcommand=v_scroll.set,
                                        xscrollcommand=h_scroll.set,
                                        highlightthickness=0,
                                        cursor="arrow")

        v_scroll.config(command=self.viewer_canvas.yview)
        h_scroll.config(command=self.viewer_canvas.xview)

        h_scroll.pack(side="bottom", fill="x")
        v_scroll.pack(side="right", fill="y")
        self.viewer_canvas.pack(fill="both", expand=True)

        # 마우스 이벤트
        self.viewer_canvas.bind("<MouseWheel>",    self._on_mousewheel)
        self.viewer_canvas.bind("<Button-4>",      self._on_mousewheel)
        self.viewer_canvas.bind("<Button-5>",      self._on_mousewheel)
        self.viewer_canvas.bind("<Button-1>",      self._on_canvas_click)
        self.viewer_canvas.bind("<B1-Motion>",     self._on_canvas_drag)
        self.viewer_canvas.bind("<ButtonRelease-1>", self._on_canvas_release)
        self.viewer_canvas.bind("<Escape>",        self._cancel_inline_entry)

        # 초기 힌트
        self.viewer_canvas.create_text(400, 300,
            text="PDF 파일을 열려면 '열기' 버튼을 클릭하세요",
            fill="#585b70", font=("Helvetica", 14), tags="hint", justify="center")

        # 기본 도구 활성화
        self._set_tool("select")

    def _build_edit_tab(self):
        tab = ttk.Frame(self.notebook)
        self.notebook.add(tab, text="  편집  ")

        # 페이지 작업
        section1 = tk.LabelFrame(tab, text=" 페이지 관리 ",
                                   bg="#1e1e2e", fg="#89b4fa",
                                   font=("Helvetica", 11, "bold"))
        section1.pack(fill="x", padx=15, pady=10)

        btn_grid = tk.Frame(section1, bg="#1e1e2e")
        btn_grid.pack(padx=10, pady=10)

        edit_buttons = [
            ("페이지 회전 (90°)", self.rotate_page),
            ("페이지 삭제", self.delete_page),
            ("페이지 복제", self.duplicate_page),
            ("페이지 이동", self.move_page),
        ]
        for i, (text, cmd) in enumerate(edit_buttons):
            tk.Button(btn_grid, text=text, command=cmd,
                      bg="#313244", fg="#cdd6f4",
                      font=("Helvetica", 10),
                      relief="flat", padx=15, pady=6,
                      cursor="hand2", width=18).grid(
                          row=i//2, column=i%2, padx=5, pady=4)

        # 텍스트 추가
        section2 = tk.LabelFrame(tab, text=" 텍스트 추가 ",
                                   bg="#1e1e2e", fg="#89b4fa",
                                   font=("Helvetica", 11, "bold"))
        section2.pack(fill="x", padx=15, pady=10)

        text_frame = tk.Frame(section2, bg="#1e1e2e")
        text_frame.pack(padx=10, pady=10, fill="x")

        tk.Label(text_frame, text="내용:", bg="#1e1e2e", fg="#cdd6f4").grid(
            row=0, column=0, sticky="w", pady=3)
        self.add_text_var = tk.StringVar()
        tk.Entry(text_frame, textvariable=self.add_text_var,
                 bg="#313244", fg="#cdd6f4", insertbackground="#cdd6f4",
                 font=("Helvetica", 11), width=40).grid(
                     row=0, column=1, padx=5, pady=3)

        tk.Label(text_frame, text="X 위치:", bg="#1e1e2e", fg="#cdd6f4").grid(
            row=1, column=0, sticky="w", pady=3)
        self.text_x_var = tk.StringVar(value="72")
        tk.Entry(text_frame, textvariable=self.text_x_var,
                 bg="#313244", fg="#cdd6f4", insertbackground="#cdd6f4",
                 width=10).grid(row=1, column=1, sticky="w", padx=5, pady=3)

        tk.Label(text_frame, text="Y 위치:", bg="#1e1e2e", fg="#cdd6f4").grid(
            row=2, column=0, sticky="w", pady=3)
        self.text_y_var = tk.StringVar(value="72")
        tk.Entry(text_frame, textvariable=self.text_y_var,
                 bg="#313244", fg="#cdd6f4", insertbackground="#cdd6f4",
                 width=10).grid(row=2, column=1, sticky="w", padx=5, pady=3)

        tk.Label(text_frame, text="크기:", bg="#1e1e2e", fg="#cdd6f4").grid(
            row=3, column=0, sticky="w", pady=3)
        self.text_size_var = tk.StringVar(value="12")
        tk.Entry(text_frame, textvariable=self.text_size_var,
                 bg="#313244", fg="#cdd6f4", insertbackground="#cdd6f4",
                 width=10).grid(row=3, column=1, sticky="w", padx=5, pady=3)

        tk.Button(text_frame, text="텍스트 추가", command=self.add_text,
                  bg="#89b4fa", fg="#1e1e2e",
                  font=("Helvetica", 10, "bold"),
                  relief="flat", padx=15, pady=6,
                  cursor="hand2").grid(row=4, column=1, sticky="w", padx=5, pady=8)

        # 이미지 추가
        section3 = tk.LabelFrame(tab, text=" 이미지 추가 ",
                                   bg="#1e1e2e", fg="#89b4fa",
                                   font=("Helvetica", 11, "bold"))
        section3.pack(fill="x", padx=15, pady=10)

        img_frame = tk.Frame(section3, bg="#1e1e2e")
        img_frame.pack(padx=10, pady=10, fill="x")

        self.img_path_var = tk.StringVar(value="이미지 파일을 선택하세요")
        tk.Entry(img_frame, textvariable=self.img_path_var,
                 bg="#313244", fg="#6c7086",
                 font=("Helvetica", 10), width=35, state="readonly").pack(
                     side="left", padx=5)
        tk.Button(img_frame, text="찾아보기", command=self.browse_image,
                  bg="#313244", fg="#cdd6f4",
                  relief="flat", padx=8, pady=4, cursor="hand2").pack(side="left", padx=3)
        tk.Button(img_frame, text="이미지 삽입", command=self.insert_image,
                  bg="#a6e3a1", fg="#1e1e2e",
                  font=("Helvetica", 10, "bold"),
                  relief="flat", padx=10, pady=4, cursor="hand2").pack(side="left", padx=3)

        # 텍스트 추출
        section4 = tk.LabelFrame(tab, text=" 텍스트 추출 ",
                                   bg="#1e1e2e", fg="#89b4fa",
                                   font=("Helvetica", 11, "bold"))
        section4.pack(fill="both", expand=True, padx=15, pady=10)

        ext_frame = tk.Frame(section4, bg="#1e1e2e")
        ext_frame.pack(fill="both", expand=True, padx=10, pady=10)

        btn_row = tk.Frame(ext_frame, bg="#1e1e2e")
        btn_row.pack(fill="x", pady=(0, 5))

        tk.Button(btn_row, text="현재 페이지 추출", command=self.extract_text_page,
                  bg="#313244", fg="#cdd6f4",
                  relief="flat", padx=10, pady=4, cursor="hand2").pack(side="left", padx=3)
        tk.Button(btn_row, text="전체 페이지 추출", command=self.extract_text_all,
                  bg="#313244", fg="#cdd6f4",
                  relief="flat", padx=10, pady=4, cursor="hand2").pack(side="left", padx=3)
        tk.Button(btn_row, text="텍스트 파일로 저장", command=self.save_text,
                  bg="#a6e3a1", fg="#1e1e2e",
                  font=("Helvetica", 10, "bold"),
                  relief="flat", padx=10, pady=4, cursor="hand2").pack(side="left", padx=3)

        text_scroll = tk.Scrollbar(ext_frame)
        self.extract_text = tk.Text(ext_frame, bg="#181825", fg="#cdd6f4",
                                     font=("Courier", 10),
                                     yscrollcommand=text_scroll.set,
                                     insertbackground="#cdd6f4",
                                     relief="flat", padx=8, pady=8)
        text_scroll.config(command=self.extract_text.yview)
        text_scroll.pack(side="right", fill="y")
        self.extract_text.pack(fill="both", expand=True)

    def _build_convert_tab(self):
        tab = ttk.Frame(self.notebook)
        self.notebook.add(tab, text="  변환  ")

        # PDF → 이미지
        s1 = tk.LabelFrame(tab, text=" PDF → 이미지 ",
                            bg="#1e1e2e", fg="#89b4fa",
                            font=("Helvetica", 11, "bold"))
        s1.pack(fill="x", padx=15, pady=10)

        f1 = tk.Frame(s1, bg="#1e1e2e")
        f1.pack(padx=10, pady=10, fill="x")

        tk.Label(f1, text="형식:", bg="#1e1e2e", fg="#cdd6f4").pack(side="left")
        self.img_format_var = tk.StringVar(value="PNG")
        ttk.Combobox(f1, textvariable=self.img_format_var,
                     values=["PNG", "JPEG", "BMP", "TIFF"],
                     width=8, state="readonly").pack(side="left", padx=5)

        tk.Label(f1, text="DPI:", bg="#1e1e2e", fg="#cdd6f4").pack(side="left")
        self.dpi_var = tk.StringVar(value="150")
        ttk.Combobox(f1, textvariable=self.dpi_var,
                     values=["72", "96", "150", "200", "300", "600"],
                     width=6, state="readonly").pack(side="left", padx=5)

        tk.Label(f1, text="페이지:", bg="#1e1e2e", fg="#cdd6f4").pack(side="left")
        self.conv_page_var = tk.StringVar(value="전체")
        ttk.Combobox(f1, textvariable=self.conv_page_var,
                     values=["전체", "현재 페이지"],
                     width=12, state="readonly").pack(side="left", padx=5)

        tk.Button(f1, text="변환 시작", command=self.pdf_to_images,
                  bg="#89b4fa", fg="#1e1e2e",
                  font=("Helvetica", 10, "bold"),
                  relief="flat", padx=15, pady=5, cursor="hand2").pack(side="left", padx=10)

        # PDF → Word
        s2 = tk.LabelFrame(tab, text=" PDF → Word (텍스트 추출) ",
                            bg="#1e1e2e", fg="#89b4fa",
                            font=("Helvetica", 11, "bold"))
        s2.pack(fill="x", padx=15, pady=10)

        f2 = tk.Frame(s2, bg="#1e1e2e")
        f2.pack(padx=10, pady=10, fill="x")

        tk.Label(f2, text="PDF의 텍스트를 Word 문서(.docx)로 추출합니다",
                 bg="#1e1e2e", fg="#a6adc8").pack(side="left")
        tk.Button(f2, text="Word로 변환", command=self.pdf_to_word,
                  bg="#89b4fa", fg="#1e1e2e",
                  font=("Helvetica", 10, "bold"),
                  relief="flat", padx=15, pady=5, cursor="hand2").pack(side="right", padx=10)

        # PDF → 텍스트
        s3 = tk.LabelFrame(tab, text=" PDF → 텍스트 ",
                            bg="#1e1e2e", fg="#89b4fa",
                            font=("Helvetica", 11, "bold"))
        s3.pack(fill="x", padx=15, pady=10)

        f3 = tk.Frame(s3, bg="#1e1e2e")
        f3.pack(padx=10, pady=10, fill="x")

        tk.Label(f3, text="PDF의 텍스트를 .txt 파일로 저장합니다",
                 bg="#1e1e2e", fg="#a6adc8").pack(side="left")
        tk.Button(f3, text="텍스트로 변환", command=self.pdf_to_text_file,
                  bg="#89b4fa", fg="#1e1e2e",
                  font=("Helvetica", 10, "bold"),
                  relief="flat", padx=15, pady=5, cursor="hand2").pack(side="right", padx=10)

        # 이미지 → PDF
        s4 = tk.LabelFrame(tab, text=" 이미지 → PDF ",
                            bg="#1e1e2e", fg="#89b4fa",
                            font=("Helvetica", 11, "bold"))
        s4.pack(fill="x", padx=15, pady=10)

        f4 = tk.Frame(s4, bg="#1e1e2e")
        f4.pack(padx=10, pady=10, fill="x")

        self.img_list_var = tk.StringVar(value="")
        self.img_listbox = tk.Listbox(f4, bg="#181825", fg="#cdd6f4",
                                       selectmode="extended",
                                       height=4, relief="flat",
                                       font=("Helvetica", 9))
        self.img_listbox.pack(side="left", fill="x", expand=True)

        btn_col = tk.Frame(f4, bg="#1e1e2e")
        btn_col.pack(side="left", padx=5)

        tk.Button(btn_col, text="추가", command=self.add_images_for_pdf,
                  bg="#313244", fg="#cdd6f4",
                  relief="flat", padx=8, pady=3, cursor="hand2").pack(pady=2)
        tk.Button(btn_col, text="삭제", command=self.remove_image_from_list,
                  bg="#313244", fg="#cdd6f4",
                  relief="flat", padx=8, pady=3, cursor="hand2").pack(pady=2)
        tk.Button(btn_col, text="PDF 생성", command=self.images_to_pdf,
                  bg="#a6e3a1", fg="#1e1e2e",
                  font=("Helvetica", 10, "bold"),
                  relief="flat", padx=8, pady=3, cursor="hand2").pack(pady=2)

        # 진행률
        prog_frame = tk.Frame(tab, bg="#1e1e2e")
        prog_frame.pack(fill="x", padx=15, pady=5)

        self.progress_var = tk.DoubleVar()
        self.progress = ttk.Progressbar(prog_frame, variable=self.progress_var,
                                         maximum=100, length=400)
        self.progress.pack(side="left", fill="x", expand=True, padx=(0, 10))
        self.progress_label = tk.Label(prog_frame, text="",
                                        bg="#1e1e2e", fg="#a6e3a1",
                                        font=("Helvetica", 9))
        self.progress_label.pack(side="left")

    def _build_merge_split_tab(self):
        tab = ttk.Frame(self.notebook)
        self.notebook.add(tab, text="  병합/분할  ")

        # PDF 병합
        s1 = tk.LabelFrame(tab, text=" PDF 병합 ",
                            bg="#1e1e2e", fg="#89b4fa",
                            font=("Helvetica", 11, "bold"))
        s1.pack(fill="x", padx=15, pady=10)

        f1 = tk.Frame(s1, bg="#1e1e2e")
        f1.pack(padx=10, pady=10, fill="x")

        self.merge_listbox = tk.Listbox(f1, bg="#181825", fg="#cdd6f4",
                                         selectmode="extended",
                                         height=5, relief="flat",
                                         font=("Helvetica", 9))
        self.merge_listbox.pack(side="left", fill="x", expand=True)

        btn_col = tk.Frame(f1, bg="#1e1e2e")
        btn_col.pack(side="left", padx=5)

        tk.Button(btn_col, text="추가", command=self.add_pdf_to_merge,
                  bg="#313244", fg="#cdd6f4",
                  relief="flat", padx=10, pady=4, cursor="hand2").pack(pady=2, fill="x")
        tk.Button(btn_col, text="삭제", command=self.remove_pdf_from_merge,
                  bg="#313244", fg="#cdd6f4",
                  relief="flat", padx=10, pady=4, cursor="hand2").pack(pady=2, fill="x")
        tk.Button(btn_col, text="위로", command=self.move_merge_up,
                  bg="#313244", fg="#cdd6f4",
                  relief="flat", padx=10, pady=4, cursor="hand2").pack(pady=2, fill="x")
        tk.Button(btn_col, text="아래로", command=self.move_merge_down,
                  bg="#313244", fg="#cdd6f4",
                  relief="flat", padx=10, pady=4, cursor="hand2").pack(pady=2, fill="x")
        tk.Button(btn_col, text="병합하기",
                  command=self.merge_pdfs,
                  bg="#a6e3a1", fg="#1e1e2e",
                  font=("Helvetica", 10, "bold"),
                  relief="flat", padx=10, pady=6, cursor="hand2").pack(pady=5, fill="x")

        # PDF 분할
        s2 = tk.LabelFrame(tab, text=" PDF 분할 ",
                            bg="#1e1e2e", fg="#89b4fa",
                            font=("Helvetica", 11, "bold"))
        s2.pack(fill="x", padx=15, pady=10)

        f2 = tk.Frame(s2, bg="#1e1e2e")
        f2.pack(padx=10, pady=10, fill="x")

        # 분할 방식
        split_opts = tk.Frame(f2, bg="#1e1e2e")
        split_opts.pack(fill="x", pady=5)

        self.split_mode = tk.StringVar(value="range")

        tk.Radiobutton(split_opts, text="페이지 범위",
                       variable=self.split_mode, value="range",
                       bg="#1e1e2e", fg="#cdd6f4",
                       selectcolor="#313244",
                       activebackground="#1e1e2e").pack(side="left", padx=5)
        tk.Radiobutton(split_opts, text="각 페이지별",
                       variable=self.split_mode, value="each",
                       bg="#1e1e2e", fg="#cdd6f4",
                       selectcolor="#313244",
                       activebackground="#1e1e2e").pack(side="left", padx=5)
        tk.Radiobutton(split_opts, text="N페이지마다",
                       variable=self.split_mode, value="every",
                       bg="#1e1e2e", fg="#cdd6f4",
                       selectcolor="#313244",
                       activebackground="#1e1e2e").pack(side="left", padx=5)

        range_frame = tk.Frame(f2, bg="#1e1e2e")
        range_frame.pack(fill="x", pady=5)

        tk.Label(range_frame, text="범위 (예: 1-3, 5, 7-10):",
                 bg="#1e1e2e", fg="#cdd6f4").pack(side="left")
        self.split_range_var = tk.StringVar()
        tk.Entry(range_frame, textvariable=self.split_range_var,
                 bg="#313244", fg="#cdd6f4", insertbackground="#cdd6f4",
                 width=20).pack(side="left", padx=5)

        tk.Label(range_frame, text="N값:",
                 bg="#1e1e2e", fg="#cdd6f4").pack(side="left", padx=(15, 0))
        self.split_n_var = tk.StringVar(value="1")
        tk.Entry(range_frame, textvariable=self.split_n_var,
                 bg="#313244", fg="#cdd6f4", insertbackground="#cdd6f4",
                 width=5).pack(side="left", padx=5)

        tk.Button(f2, text="분할하기", command=self.split_pdf,
                  bg="#a6e3a1", fg="#1e1e2e",
                  font=("Helvetica", 10, "bold"),
                  relief="flat", padx=20, pady=6, cursor="hand2").pack(anchor="w", pady=8)

    def _build_watermark_tab(self):
        tab = ttk.Frame(self.notebook)
        self.notebook.add(tab, text="  워터마크/보안  ")

        # 워터마크
        s1 = tk.LabelFrame(tab, text=" 텍스트 워터마크 ",
                            bg="#1e1e2e", fg="#89b4fa",
                            font=("Helvetica", 11, "bold"))
        s1.pack(fill="x", padx=15, pady=10)

        f1 = tk.Frame(s1, bg="#1e1e2e")
        f1.pack(padx=10, pady=10, fill="x")

        opts = [
            ("워터마크 텍스트:", "wm_text", "CONFIDENTIAL", 30),
            ("폰트 크기:", "wm_size", "40", 8),
            ("투명도 (0-1):", "wm_alpha", "0.3", 8),
            ("회전각도 (도):", "wm_angle", "45", 8),
        ]
        for row, (label, attr, default, width) in enumerate(opts):
            tk.Label(f1, text=label, bg="#1e1e2e", fg="#cdd6f4").grid(
                row=row, column=0, sticky="w", pady=3, padx=(0, 10))
            var = tk.StringVar(value=default)
            setattr(self, attr + "_var", var)
            tk.Entry(f1, textvariable=var,
                     bg="#313244", fg="#cdd6f4", insertbackground="#cdd6f4",
                     width=width).grid(row=row, column=1, sticky="w", pady=3)

        color_frame = tk.Frame(f1, bg="#1e1e2e")
        color_frame.grid(row=4, column=0, columnspan=2, sticky="w", pady=3)
        tk.Label(color_frame, text="색상:", bg="#1e1e2e", fg="#cdd6f4").pack(side="left")
        self.wm_color_var = tk.StringVar(value="회색")
        ttk.Combobox(color_frame, textvariable=self.wm_color_var,
                     values=["회색", "빨강", "파랑", "초록", "검정"],
                     width=8, state="readonly").pack(side="left", padx=5)

        self.wm_all_pages = tk.BooleanVar(value=True)
        tk.Checkbutton(f1, text="전체 페이지에 적용",
                       variable=self.wm_all_pages,
                       bg="#1e1e2e", fg="#cdd6f4",
                       selectcolor="#313244",
                       activebackground="#1e1e2e").grid(
                           row=5, column=0, columnspan=2, sticky="w", pady=3)

        tk.Button(f1, text="워터마크 추가", command=self.add_watermark,
                  bg="#89b4fa", fg="#1e1e2e",
                  font=("Helvetica", 10, "bold"),
                  relief="flat", padx=15, pady=6, cursor="hand2").grid(
                      row=6, column=1, sticky="w", pady=8)

        # 보안
        s2 = tk.LabelFrame(tab, text=" 비밀번호 보호 ",
                            bg="#1e1e2e", fg="#89b4fa",
                            font=("Helvetica", 11, "bold"))
        s2.pack(fill="x", padx=15, pady=10)

        f2 = tk.Frame(s2, bg="#1e1e2e")
        f2.pack(padx=10, pady=10, fill="x")

        tk.Label(f2, text="비밀번호:", bg="#1e1e2e", fg="#cdd6f4").grid(
            row=0, column=0, sticky="w", pady=3)
        self.pdf_password_var = tk.StringVar()
        tk.Entry(f2, textvariable=self.pdf_password_var,
                 show="*", bg="#313244", fg="#cdd6f4",
                 insertbackground="#cdd6f4", width=25).grid(
                     row=0, column=1, padx=5, pady=3, sticky="w")

        tk.Button(f2, text="비밀번호 설정", command=self.set_password,
                  bg="#f38ba8", fg="#1e1e2e",
                  font=("Helvetica", 10, "bold"),
                  relief="flat", padx=12, pady=5, cursor="hand2").grid(
                      row=1, column=1, sticky="w", padx=5, pady=5)

        # PDF 최적화
        s3 = tk.LabelFrame(tab, text=" PDF 최적화 ",
                            bg="#1e1e2e", fg="#89b4fa",
                            font=("Helvetica", 11, "bold"))
        s3.pack(fill="x", padx=15, pady=10)

        f3 = tk.Frame(s3, bg="#1e1e2e")
        f3.pack(padx=10, pady=10, fill="x")

        tk.Label(f3, text="파일 크기를 줄이고 PDF를 최적화합니다",
                 bg="#1e1e2e", fg="#a6adc8").pack(side="left")
        tk.Button(f3, text="최적화", command=self.optimize_pdf,
                  bg="#89b4fa", fg="#1e1e2e",
                  font=("Helvetica", 10, "bold"),
                  relief="flat", padx=15, pady=5, cursor="hand2").pack(side="right", padx=10)

    # ─── 파일 작업 ───────────────────────────────────────────

    def open_pdf(self):
        path = filedialog.askopenfilename(
            title="PDF 파일 열기",
            filetypes=[("PDF 파일", "*.pdf"), ("모든 파일", "*.*")])
        if not path:
            return
        self._load_pdf(path)

    def _load_pdf(self, path):
        try:
            if self.current_pdf:
                self.current_pdf.close()
            self.current_pdf = fitz.open(path)
            self.current_path = path
            self.current_page = 0
            self.total_pages = len(self.current_pdf)
            self.zoom = 1.0
            self._update_thumbnails()
            self._show_page()
            fname = os.path.basename(path)
            self.status_var.set(f"열림: {fname}  |  {self.total_pages}페이지")
            self.root.title(f"PDF 도구 - {fname}")
        except Exception as e:
            messagebox.showerror("오류", f"파일을 열 수 없습니다:\n{e}")

    def save_pdf(self):
        if not self.current_pdf:
            messagebox.showwarning("경고", "열린 PDF가 없습니다")
            return
        if not self.current_path:
            self.save_pdf_as()
            return
        try:
            self.current_pdf.save(self.current_path, garbage=4, deflate=True)
            self.status_var.set(f"저장됨: {os.path.basename(self.current_path)}")
        except Exception as e:
            messagebox.showerror("오류", f"저장 실패:\n{e}")

    def save_pdf_as(self):
        if not self.current_pdf:
            messagebox.showwarning("경고", "열린 PDF가 없습니다")
            return
        path = filedialog.asksaveasfilename(
            title="다른 이름으로 저장",
            defaultextension=".pdf",
            filetypes=[("PDF 파일", "*.pdf")])
        if not path:
            return
        try:
            self.current_pdf.save(path, garbage=4, deflate=True)
            self.current_path = path
            self.status_var.set(f"저장됨: {os.path.basename(path)}")
        except Exception as e:
            messagebox.showerror("오류", f"저장 실패:\n{e}")

    # ─── 뷰어 ────────────────────────────────────────────────

    def _show_page(self):
        if not self.current_pdf:
            return
        try:
            page = self.current_pdf[self.current_page]
            mat = fitz.Matrix(self.zoom * 1.5, self.zoom * 1.5)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            img_data = pix.tobytes("png")
            img = Image.open(io.BytesIO(img_data))
            photo = ImageTk.PhotoImage(img)

            self.viewer_canvas.delete("all")
            w, h = img.size
            self.viewer_canvas.configure(scrollregion=(0, 0, w + 40, h + 40))
            self.viewer_canvas.create_image(20, 20, anchor="nw", image=photo)
            self.photo_images = [photo]

            self.page_var.set(f"{self.current_page + 1} / {self.total_pages}")
            self.zoom_var.set(f"{int(self.zoom * 100)}%")
        except Exception as e:
            self.status_var.set(f"페이지 표시 오류: {e}")

    def _update_thumbnails(self):
        for w in self.thumb_inner.winfo_children():
            w.destroy()
        self.thumb_photos = []

        if not self.current_pdf:
            return

        for i in range(min(self.total_pages, 100)):
            try:
                page = self.current_pdf[i]
                mat = fitz.Matrix(0.2, 0.2)
                pix = page.get_pixmap(matrix=mat, alpha=False)
                img = Image.open(io.BytesIO(pix.tobytes("png")))
                photo = ImageTk.PhotoImage(img)
                self.thumb_photos.append(photo)

                frame = tk.Frame(self.thumb_inner, bg="#181825",
                                  cursor="hand2")
                frame.pack(pady=3, padx=5)

                lbl = tk.Label(frame, image=photo, bg="#313244",
                                relief="ridge", bd=1)
                lbl.pack()
                tk.Label(frame, text=str(i + 1),
                         bg="#181825", fg="#6c7086",
                         font=("Helvetica", 8)).pack()

                page_idx = i
                frame.bind("<Button-1>",
                            lambda e, p=page_idx: self._goto_page(p))
                lbl.bind("<Button-1>",
                          lambda e, p=page_idx: self._goto_page(p))
            except Exception:
                pass

    def _goto_page(self, page_idx):
        self.current_page = page_idx
        self._show_page()

    def first_page(self):
        if self.current_pdf:
            self.current_page = 0
            self._show_page()

    def last_page(self):
        if self.current_pdf:
            self.current_page = self.total_pages - 1
            self._show_page()

    def prev_page(self):
        if self.current_pdf and self.current_page > 0:
            self.current_page -= 1
            self._show_page()

    def next_page(self):
        if self.current_pdf and self.current_page < self.total_pages - 1:
            self.current_page += 1
            self._show_page()

    def zoom_in(self):
        self.zoom = min(self.zoom * 1.25, 5.0)
        self._show_page()

    def zoom_out(self):
        self.zoom = max(self.zoom / 1.25, 0.2)
        self._show_page()

    def zoom_fit(self):
        self.zoom = 1.0
        self._show_page()

    def _on_mousewheel(self, event):
        if self.edit_tool == "select":
            if event.num == 4 or event.delta > 0:
                self.viewer_canvas.yview_scroll(-1, "units")
            else:
                self.viewer_canvas.yview_scroll(1, "units")

    # ─── 인라인 편집 도구 ─────────────────────────────────────

    def _set_tool(self, tool_id):
        self.edit_tool = tool_id
        cursors = {
            "select":    "arrow",
            "text":      "xterm",
            "highlight": "crosshair",
            "rect":      "crosshair",
            "pencil":    "pencil",
            "arrow":     "crosshair",
            "redact":    "crosshair",
        }
        hints = {
            "select":    "↖ 선택 — 스크롤로 페이지 이동",
            "text":      "T 텍스트 — 클릭한 위치에 텍스트 입력",
            "highlight": "▬ 형광펜 — 드래그하여 영역 강조",
            "rect":      "▭ 사각형 — 드래그하여 사각형 그리기",
            "pencil":    "✏ 펜 — 자유롭게 그리기",
            "arrow":     "→ 화살표 — 드래그하여 화살표 그리기",
            "redact":    "█ 가리기 — 드래그하여 내용 가리기",
        }
        self.viewer_canvas.configure(cursor=cursors.get(tool_id, "arrow"))
        self.tool_hint_var.set(hints.get(tool_id, ""))

        for tid, btn in self.tool_btns.items():
            if tid == tool_id:
                btn.configure(relief="sunken", bg="#45475a")
            else:
                btn.configure(relief="flat", bg="#313244")

    def _set_color(self, hex_c, rgb_c):
        self.draw_color_hex = hex_c
        self.draw_color = rgb_c
        self.color_indicator.configure(bg=hex_c)

    def _canvas_to_pdf(self, cx, cy):
        """캔버스 픽셀 좌표 → PDF 포인트 좌표"""
        scale = self.zoom * 1.5
        pdf_x = (cx - 20) / scale
        pdf_y = (cy - 20) / scale
        return pdf_x, pdf_y

    def _save_undo(self):
        if not self.current_pdf:
            return
        buf = io.BytesIO()
        self.current_pdf.save(buf)
        self.undo_stack.append(buf.getvalue())
        if len(self.undo_stack) > 20:
            self.undo_stack.pop(0)

    def undo(self):
        if not self.undo_stack:
            self.status_var.set("실행취소할 내용이 없습니다")
            return
        data = self.undo_stack.pop()
        self.current_pdf = fitz.open("pdf", data)
        self.total_pages = len(self.current_pdf)
        self._show_page()
        self._update_thumbnails()
        self.status_var.set("실행취소됨")

    # ── 마우스 이벤트 핸들러 ───────────────────────────────────

    def _on_canvas_click(self, event):
        if not self.current_pdf:
            return

        # 인라인 입력창이 열려 있으면 먼저 확정
        if self.inline_entry_win:
            self._commit_inline_entry()

        cx = self.viewer_canvas.canvasx(event.x)
        cy = self.viewer_canvas.canvasy(event.y)
        self.draw_start = (cx, cy)
        self.freehand_points = [(cx, cy)]
        self.freehand_canvas_ids = []

        if self.edit_tool == "text":
            self._open_inline_entry(cx, cy)

    def _on_canvas_drag(self, event):
        if not self.current_pdf or not self.draw_start:
            return

        cx = self.viewer_canvas.canvasx(event.x)
        cy = self.viewer_canvas.canvasy(event.y)
        x0, y0 = self.draw_start

        if self.edit_tool == "pencil":
            self.freehand_points.append((cx, cy))
            if len(self.freehand_points) >= 2:
                p1 = self.freehand_points[-2]
                p2 = self.freehand_points[-1]
                cid = self.viewer_canvas.create_line(
                    p1[0], p1[1], p2[0], p2[1],
                    fill=self.draw_color_hex,
                    width=self.draw_size, smooth=True)
                self.freehand_canvas_ids.append(cid)

        elif self.edit_tool in ("highlight", "rect", "redact", "arrow"):
            if self.draw_preview:
                self.viewer_canvas.delete(self.draw_preview)
            if self.edit_tool == "arrow":
                self.draw_preview = self.viewer_canvas.create_line(
                    x0, y0, cx, cy,
                    fill=self.draw_color_hex,
                    width=max(self.draw_size, 2),
                    arrow="last")
            elif self.edit_tool == "highlight":
                self.draw_preview = self.viewer_canvas.create_rectangle(
                    x0, y0, cx, cy,
                    fill=self.draw_color_hex,
                    outline="", stipple="gray50")
            elif self.edit_tool == "redact":
                self.draw_preview = self.viewer_canvas.create_rectangle(
                    x0, y0, cx, cy,
                    fill="#000000", outline="")
            else:
                self.draw_preview = self.viewer_canvas.create_rectangle(
                    x0, y0, cx, cy,
                    outline=self.draw_color_hex,
                    width=self.draw_size)

    def _on_canvas_release(self, event):
        if not self.current_pdf or not self.draw_start:
            return

        cx = self.viewer_canvas.canvasx(event.x)
        cy = self.viewer_canvas.canvasy(event.y)
        x0, y0 = self.draw_start

        tool = self.edit_tool

        if tool == "text":
            pass  # 클릭 시 이미 처리됨

        elif tool == "pencil":
            if len(self.freehand_points) > 1:
                self._save_undo()
                page = self.current_pdf[self.current_page]
                pts = [self._canvas_to_pdf(p[0], p[1])
                       for p in self.freehand_points]
                # 폴리라인을 짧은 선분들로 추가
                for i in range(len(pts) - 1):
                    p1 = fitz.Point(pts[i])
                    p2 = fitz.Point(pts[i+1])
                    annot = page.add_line_annot(p1, p2)
                    annot.set_colors(stroke=self.draw_color)
                    annot.set_border(width=self.draw_size)
                    annot.update()
                self._show_page()
                self.status_var.set("펜 그리기 완료")
            for cid in self.freehand_canvas_ids:
                self.viewer_canvas.delete(cid)
            self.freehand_canvas_ids = []

        elif tool in ("highlight", "rect", "redact", "arrow"):
            if self.draw_preview:
                self.viewer_canvas.delete(self.draw_preview)
                self.draw_preview = None

            if abs(cx - x0) < 3 and abs(cy - y0) < 3:
                self.draw_start = None
                return

            self._save_undo()
            page = self.current_pdf[self.current_page]
            px0, py0 = self._canvas_to_pdf(x0, y0)
            px1, py1 = self._canvas_to_pdf(cx, cy)
            rect = fitz.Rect(
                min(px0, px1), min(py0, py1),
                max(px0, px1), max(py0, py1))

            if tool == "highlight":
                annot = page.add_highlight_annot(rect)
                r, g, b = self.draw_color
                annot.set_colors(stroke=(r, g, b))
                annot.update()
            elif tool == "rect":
                annot = page.add_rect_annot(rect)
                annot.set_colors(stroke=self.draw_color, fill=None)
                annot.set_border(width=self.draw_size)
                annot.update()
            elif tool == "redact":
                page.add_redact_annot(rect, fill=(0, 0, 0))
                page.apply_redactions()
            elif tool == "arrow":
                p1 = fitz.Point(px0, py0)
                p2 = fitz.Point(px1, py1)
                annot = page.add_line_annot(p1, p2)
                annot.set_colors(stroke=self.draw_color)
                annot.set_border(width=max(self.draw_size, 1))
                annot.update(line_end_style="OpenArrow")

            self._show_page()
            self._update_thumbnails()
            self.status_var.set(f"{tool} 추가됨")

        self.draw_start = None

    # ── 인라인 텍스트 입력 ─────────────────────────────────────

    def _open_inline_entry(self, cx, cy):
        """캔버스 위에 직접 텍스트 입력창 표시"""
        self.inline_entry_pos = (cx, cy)

        frame = tk.Frame(self.viewer_canvas,
                         bg="#1e1e2e", bd=1, relief="solid",
                         highlightbackground=self.draw_color_hex,
                         highlightthickness=2)

        entry = tk.Entry(frame,
                         bg="#1e1e2e", fg=self.draw_color_hex,
                         insertbackground=self.draw_color_hex,
                         font=("Helvetica", max(self.text_size, 9)),
                         bd=0, width=24,
                         relief="flat")
        entry.pack(padx=4, pady=2)

        btn_row = tk.Frame(frame, bg="#1e1e2e")
        btn_row.pack(fill="x")
        tk.Button(btn_row, text="확인", command=self._commit_inline_entry,
                  bg=self.draw_color_hex, fg="#1e1e2e",
                  font=("Helvetica", 8, "bold"), relief="flat",
                  cursor="hand2").pack(side="left", padx=2, pady=2)
        tk.Button(btn_row, text="취소", command=self._cancel_inline_entry,
                  bg="#313244", fg="#cdd6f4",
                  font=("Helvetica", 8), relief="flat",
                  cursor="hand2").pack(side="left", padx=2, pady=2)

        hint = tk.Label(frame, text="Enter=확인  Esc=취소",
                        bg="#1e1e2e", fg="#585b70",
                        font=("Helvetica", 7))
        hint.pack(anchor="w", padx=4)

        win = self.viewer_canvas.create_window(cx, cy, anchor="nw", window=frame)
        self.inline_entry_win = win
        self.inline_entry_widget = entry

        entry.focus_set()
        entry.bind("<Return>", lambda e: self._commit_inline_entry())
        entry.bind("<Escape>", lambda e: self._cancel_inline_entry())

    def _commit_inline_entry(self):
        if not self.inline_entry_win:
            return
        text = self.inline_entry_widget.get().strip()
        cx, cy = self.inline_entry_pos
        self.viewer_canvas.delete(self.inline_entry_win)
        self.inline_entry_win = None

        if text and self.current_pdf:
            self._save_undo()
            pdf_x, pdf_y = self._canvas_to_pdf(cx, cy)
            page = self.current_pdf[self.current_page]
            # 텍스트 주석으로 추가 (FreeText)
            fs = self.text_size
            w = max(len(text) * fs * 0.65, 50)
            rect = fitz.Rect(pdf_x, pdf_y - fs * 1.2, pdf_x + w, pdf_y + fs * 0.5)
            annot = page.add_freetext_annot(
                rect, text,
                fontsize=fs,
                text_color=self.draw_color,
                fill_color=None,
                align=0,
            )
            annot.update()
            self._show_page()
            self._update_thumbnails()
            self.status_var.set(f"텍스트 추가됨: \"{text}\"")

    def _cancel_inline_entry(self, event=None):
        if self.inline_entry_win:
            self.viewer_canvas.delete(self.inline_entry_win)
            self.inline_entry_win = None

    # ─── 편집 ────────────────────────────────────────────────

    def _check_pdf(self):
        if not self.current_pdf:
            messagebox.showwarning("경고", "PDF 파일을 먼저 열어주세요")
            return False
        return True

    def rotate_page(self):
        if not self._check_pdf():
            return
        page = self.current_pdf[self.current_page]
        page.set_rotation((page.rotation + 90) % 360)
        self._show_page()
        self._update_thumbnails()
        self.status_var.set(f"페이지 {self.current_page + 1} 회전됨")

    def delete_page(self):
        if not self._check_pdf():
            return
        if self.total_pages <= 1:
            messagebox.showwarning("경고", "마지막 페이지는 삭제할 수 없습니다")
            return
        if messagebox.askyesno("확인", f"페이지 {self.current_page + 1}을 삭제하시겠습니까?"):
            self.current_pdf.delete_page(self.current_page)
            self.total_pages -= 1
            self.current_page = min(self.current_page, self.total_pages - 1)
            self._update_thumbnails()
            self._show_page()
            self.status_var.set(f"페이지 삭제됨  |  현재 {self.total_pages}페이지")

    def duplicate_page(self):
        if not self._check_pdf():
            return
        self.current_pdf.copy_page(self.current_page, self.current_page + 1)
        self.total_pages += 1
        self._update_thumbnails()
        self._show_page()
        self.status_var.set(f"페이지 복제됨")

    def move_page(self):
        if not self._check_pdf():
            return
        target = simpledialog.askinteger(
            "페이지 이동",
            f"이동할 위치 (1 ~ {self.total_pages}):",
            minvalue=1, maxvalue=self.total_pages)
        if target is None:
            return
        self.current_pdf.move_page(self.current_page, target - 1)
        self.current_page = target - 1
        self._update_thumbnails()
        self._show_page()
        self.status_var.set(f"페이지를 {target}번으로 이동")

    def add_text(self):
        if not self._check_pdf():
            return
        text = self.add_text_var.get().strip()
        if not text:
            messagebox.showwarning("경고", "텍스트를 입력하세요")
            return
        try:
            x = float(self.text_x_var.get())
            y = float(self.text_y_var.get())
            size = float(self.text_size_var.get())
        except ValueError:
            messagebox.showerror("오류", "올바른 숫자를 입력하세요")
            return

        page = self.current_pdf[self.current_page]
        page.insert_text((x, y), text,
                          fontsize=size,
                          color=(0, 0, 0))
        self._show_page()
        self.status_var.set("텍스트 추가됨")

    def browse_image(self):
        path = filedialog.askopenfilename(
            title="이미지 선택",
            filetypes=[("이미지", "*.png *.jpg *.jpeg *.bmp *.gif *.tiff"),
                       ("모든 파일", "*.*")])
        if path:
            self.img_path_var.set(path)

    def insert_image(self):
        if not self._check_pdf():
            return
        img_path = self.img_path_var.get()
        if not os.path.exists(img_path):
            messagebox.showwarning("경고", "이미지 파일을 선택하세요")
            return
        try:
            page = self.current_pdf[self.current_page]
            rect = fitz.Rect(50, 50, 300, 300)
            page.insert_image(rect, filename=img_path)
            self._show_page()
            self.status_var.set("이미지 삽입됨")
        except Exception as e:
            messagebox.showerror("오류", f"이미지 삽입 실패:\n{e}")

    def extract_text_page(self):
        if not self._check_pdf():
            return
        page = self.current_pdf[self.current_page]
        text = page.get_text()
        self.extract_text.delete("1.0", "end")
        self.extract_text.insert("1.0", text)
        self.status_var.set(f"페이지 {self.current_page + 1} 텍스트 추출됨")

    def extract_text_all(self):
        if not self._check_pdf():
            return
        all_text = []
        for i, page in enumerate(self.current_pdf):
            all_text.append(f"=== 페이지 {i + 1} ===\n")
            all_text.append(page.get_text())
            all_text.append("\n")
        self.extract_text.delete("1.0", "end")
        self.extract_text.insert("1.0", "".join(all_text))
        self.status_var.set("전체 텍스트 추출됨")

    def save_text(self):
        text = self.extract_text.get("1.0", "end").strip()
        if not text:
            messagebox.showwarning("경고", "추출된 텍스트가 없습니다")
            return
        path = filedialog.asksaveasfilename(
            title="텍스트 저장",
            defaultextension=".txt",
            filetypes=[("텍스트 파일", "*.txt")])
        if path:
            with open(path, "w", encoding="utf-8") as f:
                f.write(text)
            self.status_var.set(f"텍스트 저장됨: {os.path.basename(path)}")

    # ─── 변환 ────────────────────────────────────────────────

    def pdf_to_images(self):
        if not self._check_pdf():
            return

        save_dir = filedialog.askdirectory(title="저장 폴더 선택")
        if not save_dir:
            return

        fmt = self.img_format_var.get().lower()
        dpi = int(self.dpi_var.get())
        mat = fitz.Matrix(dpi / 72, dpi / 72)
        mode = self.conv_page_var.get()

        pages = [self.current_page] if mode == "현재 페이지" else range(self.total_pages)
        total = len(list(pages))
        pages = [self.current_page] if mode == "현재 페이지" else range(self.total_pages)

        def do_convert():
            for idx, page_num in enumerate(pages):
                page = self.current_pdf[page_num]
                pix = page.get_pixmap(matrix=mat, alpha=False)
                fname = os.path.basename(self.current_path or "output")
                fname = os.path.splitext(fname)[0]
                out_path = os.path.join(save_dir, f"{fname}_page{page_num + 1}.{fmt}")

                if fmt == "jpeg" or fmt == "jpg":
                    img = Image.open(io.BytesIO(pix.tobytes("png")))
                    img.save(out_path, "JPEG", quality=95)
                else:
                    img = Image.open(io.BytesIO(pix.tobytes("png")))
                    img.save(out_path, fmt.upper())

                progress = (idx + 1) / total * 100
                self.progress_var.set(progress)
                self.progress_label.config(
                    text=f"{idx + 1}/{total} 변환 중...")
                self.root.update_idletasks()

            self.progress_label.config(text=f"완료! {total}개 이미지 저장됨")
            self.status_var.set(f"이미지 변환 완료: {save_dir}")
            messagebox.showinfo("완료",
                f"{total}개 이미지가 저장되었습니다\n{save_dir}")

        threading.Thread(target=do_convert, daemon=True).start()

    def pdf_to_word(self):
        if not self._check_pdf():
            return

        path = filedialog.asksaveasfilename(
            title="Word 파일 저장",
            defaultextension=".docx",
            filetypes=[("Word 문서", "*.docx")])
        if not path:
            return

        try:
            doc = Document()
            doc.add_heading(os.path.basename(self.current_path or "PDF"), 0)

            for i, page in enumerate(self.current_pdf):
                doc.add_heading(f"페이지 {i + 1}", 2)
                text = page.get_text()
                if text.strip():
                    doc.add_paragraph(text)
                else:
                    doc.add_paragraph("(텍스트 없음)")

                self.progress_var.set((i + 1) / self.total_pages * 100)
                self.root.update_idletasks()

            doc.save(path)
            self.progress_label.config(text="완료!")
            self.status_var.set(f"Word 변환 완료: {os.path.basename(path)}")
            messagebox.showinfo("완료", f"Word 파일로 저장되었습니다:\n{path}")
        except Exception as e:
            messagebox.showerror("오류", f"Word 변환 실패:\n{e}")

    def pdf_to_text_file(self):
        if not self._check_pdf():
            return

        path = filedialog.asksaveasfilename(
            title="텍스트 파일 저장",
            defaultextension=".txt",
            filetypes=[("텍스트 파일", "*.txt")])
        if not path:
            return

        try:
            with open(path, "w", encoding="utf-8") as f:
                for i, page in enumerate(self.current_pdf):
                    f.write(f"=== 페이지 {i + 1} ===\n")
                    f.write(page.get_text())
                    f.write("\n\n")
                    self.progress_var.set((i + 1) / self.total_pages * 100)
                    self.root.update_idletasks()

            self.progress_label.config(text="완료!")
            self.status_var.set(f"텍스트 저장됨: {os.path.basename(path)}")
            messagebox.showinfo("완료", f"텍스트 파일로 저장되었습니다:\n{path}")
        except Exception as e:
            messagebox.showerror("오류", f"변환 실패:\n{e}")

    def add_images_for_pdf(self):
        paths = filedialog.askopenfilenames(
            title="이미지 선택",
            filetypes=[("이미지", "*.png *.jpg *.jpeg *.bmp *.tiff"),
                       ("모든 파일", "*.*")])
        for path in paths:
            self.img_listbox.insert("end", path)

    def remove_image_from_list(self):
        sel = self.img_listbox.curselection()
        for idx in reversed(sel):
            self.img_listbox.delete(idx)

    def images_to_pdf(self):
        items = list(self.img_listbox.get(0, "end"))
        if not items:
            messagebox.showwarning("경고", "이미지를 추가하세요")
            return

        path = filedialog.asksaveasfilename(
            title="PDF 저장",
            defaultextension=".pdf",
            filetypes=[("PDF 파일", "*.pdf")])
        if not path:
            return

        try:
            doc = fitz.open()
            for i, img_path in enumerate(items):
                img = Image.open(img_path)
                w, h = img.size
                page = doc.new_page(width=w, height=h)
                page.insert_image(fitz.Rect(0, 0, w, h), filename=img_path)
                self.progress_var.set((i + 1) / len(items) * 100)
                self.progress_label.config(text=f"{i + 1}/{len(items)}...")
                self.root.update_idletasks()

            doc.save(path, garbage=4, deflate=True)
            doc.close()
            self.progress_label.config(text="완료!")
            self.status_var.set(f"PDF 생성 완료: {os.path.basename(path)}")
            messagebox.showinfo("완료", f"PDF가 생성되었습니다:\n{path}")
        except Exception as e:
            messagebox.showerror("오류", f"변환 실패:\n{e}")

    # ─── 병합/분할 ───────────────────────────────────────────

    def add_pdf_to_merge(self):
        paths = filedialog.askopenfilenames(
            title="PDF 파일 선택",
            filetypes=[("PDF 파일", "*.pdf")])
        for path in paths:
            self.merge_listbox.insert("end", path)

    def remove_pdf_from_merge(self):
        sel = self.merge_listbox.curselection()
        for idx in reversed(sel):
            self.merge_listbox.delete(idx)

    def move_merge_up(self):
        sel = self.merge_listbox.curselection()
        if not sel or sel[0] == 0:
            return
        idx = sel[0]
        text = self.merge_listbox.get(idx)
        self.merge_listbox.delete(idx)
        self.merge_listbox.insert(idx - 1, text)
        self.merge_listbox.select_set(idx - 1)

    def move_merge_down(self):
        sel = self.merge_listbox.curselection()
        if not sel or sel[-1] == self.merge_listbox.size() - 1:
            return
        idx = sel[-1]
        text = self.merge_listbox.get(idx)
        self.merge_listbox.delete(idx)
        self.merge_listbox.insert(idx + 1, text)
        self.merge_listbox.select_set(idx + 1)

    def merge_pdfs(self):
        items = list(self.merge_listbox.get(0, "end"))
        if len(items) < 2:
            messagebox.showwarning("경고", "PDF를 2개 이상 추가하세요")
            return

        path = filedialog.asksaveasfilename(
            title="병합된 PDF 저장",
            defaultextension=".pdf",
            filetypes=[("PDF 파일", "*.pdf")])
        if not path:
            return

        try:
            merged = fitz.open()
            for pdf_path in items:
                doc = fitz.open(pdf_path)
                merged.insert_pdf(doc)
                doc.close()
            merged.save(path, garbage=4, deflate=True)
            merged.close()
            self.status_var.set(f"병합 완료: {os.path.basename(path)}")
            messagebox.showinfo("완료",
                f"{len(items)}개 PDF가 병합되었습니다:\n{path}")
        except Exception as e:
            messagebox.showerror("오류", f"병합 실패:\n{e}")

    def split_pdf(self):
        if not self._check_pdf():
            return

        save_dir = filedialog.askdirectory(title="저장 폴더 선택")
        if not save_dir:
            return

        mode = self.split_mode.get()
        base_name = os.path.splitext(
            os.path.basename(self.current_path or "output"))[0]

        try:
            if mode == "each":
                for i in range(self.total_pages):
                    doc = fitz.open()
                    doc.insert_pdf(self.current_pdf, from_page=i, to_page=i)
                    out = os.path.join(save_dir, f"{base_name}_page{i+1}.pdf")
                    doc.save(out, garbage=4, deflate=True)
                    doc.close()
                messagebox.showinfo("완료",
                    f"{self.total_pages}개 파일로 분할됨\n{save_dir}")

            elif mode == "range":
                range_str = self.split_range_var.get().strip()
                if not range_str:
                    messagebox.showwarning("경고", "페이지 범위를 입력하세요")
                    return
                pages = self._parse_page_range(range_str)
                doc = fitz.open()
                for p in pages:
                    if 0 <= p < self.total_pages:
                        doc.insert_pdf(self.current_pdf, from_page=p, to_page=p)
                out = os.path.join(save_dir, f"{base_name}_split.pdf")
                doc.save(out, garbage=4, deflate=True)
                doc.close()
                messagebox.showinfo("완료", f"선택된 페이지 저장됨:\n{out}")

            elif mode == "every":
                n = int(self.split_n_var.get())
                if n < 1:
                    messagebox.showwarning("경고", "N값은 1 이상이어야 합니다")
                    return
                count = 0
                for start in range(0, self.total_pages, n):
                    end = min(start + n - 1, self.total_pages - 1)
                    doc = fitz.open()
                    doc.insert_pdf(self.current_pdf, from_page=start, to_page=end)
                    out = os.path.join(save_dir,
                        f"{base_name}_part{count+1}.pdf")
                    doc.save(out, garbage=4, deflate=True)
                    doc.close()
                    count += 1
                messagebox.showinfo("완료",
                    f"{count}개 파일로 분할됨\n{save_dir}")

            self.status_var.set(f"분할 완료: {save_dir}")
        except Exception as e:
            messagebox.showerror("오류", f"분할 실패:\n{e}")

    def _parse_page_range(self, range_str):
        pages = set()
        for part in range_str.split(","):
            part = part.strip()
            if "-" in part:
                a, b = part.split("-", 1)
                try:
                    for p in range(int(a.strip()) - 1, int(b.strip())):
                        pages.add(p)
                except ValueError:
                    pass
            else:
                try:
                    pages.add(int(part) - 1)
                except ValueError:
                    pass
        return sorted(pages)

    # ─── 워터마크 / 보안 ─────────────────────────────────────

    def add_watermark(self):
        if not self._check_pdf():
            return

        text = self.wm_text_var.get().strip()
        if not text:
            messagebox.showwarning("경고", "워터마크 텍스트를 입력하세요")
            return

        try:
            size = float(self.wm_size_var.get())
            alpha = float(self.wm_alpha_var.get())
            angle = float(self.wm_angle_var.get())
        except ValueError:
            messagebox.showerror("오류", "올바른 값을 입력하세요")
            return

        color_map = {
            "회색": (0.5, 0.5, 0.5),
            "빨강": (1, 0, 0),
            "파랑": (0, 0, 1),
            "초록": (0, 0.6, 0),
            "검정": (0, 0, 0),
        }
        color = color_map.get(self.wm_color_var.get(), (0.5, 0.5, 0.5))

        pages = range(self.total_pages) if self.wm_all_pages.get() else [self.current_page]

        try:
            for page_num in pages:
                page = self.current_pdf[page_num]
                rect = page.rect
                cx, cy = rect.width / 2, rect.height / 2

                page.insert_text(
                    (cx - size * len(text) * 0.3, cy),
                    text,
                    fontsize=size,
                    color=color,
                    rotate=angle,
                    fill_opacity=alpha,
                )

            self._show_page()
            self.status_var.set("워터마크 추가됨")
            messagebox.showinfo("완료", "워터마크가 추가되었습니다")
        except Exception as e:
            messagebox.showerror("오류", f"워터마크 추가 실패:\n{e}")

    def set_password(self):
        if not self._check_pdf():
            return

        password = self.pdf_password_var.get()
        if not password:
            messagebox.showwarning("경고", "비밀번호를 입력하세요")
            return

        path = filedialog.asksaveasfilename(
            title="암호화된 PDF 저장",
            defaultextension=".pdf",
            filetypes=[("PDF 파일", "*.pdf")])
        if not path:
            return

        try:
            perm = (
                fitz.PDF_PERM_ACCESSIBILITY
                | fitz.PDF_PERM_PRINT
                | fitz.PDF_PERM_COPY
                | fitz.PDF_PERM_ANNOTATE
            )
            encrypt_meth = fitz.PDF_ENCRYPT_AES_256

            self.current_pdf.save(
                path,
                encryption=encrypt_meth,
                owner_pw=password + "_owner",
                user_pw=password,
                permissions=perm,
                garbage=4,
                deflate=True,
            )
            self.status_var.set(f"암호화됨: {os.path.basename(path)}")
            messagebox.showinfo("완료", f"비밀번호가 설정되었습니다:\n{path}")
        except Exception as e:
            messagebox.showerror("오류", f"암호화 실패:\n{e}")

    def optimize_pdf(self):
        if not self._check_pdf():
            return

        path = filedialog.asksaveasfilename(
            title="최적화된 PDF 저장",
            defaultextension=".pdf",
            filetypes=[("PDF 파일", "*.pdf")])
        if not path:
            return

        try:
            orig_size = os.path.getsize(self.current_path) if self.current_path else 0
            self.current_pdf.save(
                path,
                garbage=4,
                deflate=True,
                clean=True,
                deflate_images=True,
                deflate_fonts=True,
            )
            new_size = os.path.getsize(path)
            saved = orig_size - new_size
            pct = (saved / orig_size * 100) if orig_size > 0 else 0

            self.status_var.set(f"최적화 완료: {os.path.basename(path)}")
            messagebox.showinfo("완료",
                f"PDF 최적화 완료!\n\n"
                f"원본: {orig_size/1024:.1f} KB\n"
                f"최적화: {new_size/1024:.1f} KB\n"
                f"절약: {saved/1024:.1f} KB ({pct:.1f}%)")
        except Exception as e:
            messagebox.showerror("오류", f"최적화 실패:\n{e}")


def main():
    root = tk.Tk()
    app = PDFTool(root)

    # 커맨드라인으로 파일 열기
    if len(sys.argv) > 1 and os.path.exists(sys.argv[1]):
        root.after(100, lambda: app._load_pdf(sys.argv[1]))

    root.mainloop()


if __name__ == "__main__":
    main()
